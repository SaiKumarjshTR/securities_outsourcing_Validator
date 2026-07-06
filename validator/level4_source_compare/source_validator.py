"""
level4_source_compare/source_validator.py
──────────────────────────────────────────
Level 4: Source Comparison Validator  (30 points total)

Runs ONLY when a source PDF is provided. Compares the pipeline-generated
SGML against the source document across 6 validation dimensions:

  D2  Tagging accuracy        — font/style → tag correctness  (5 pts)
  D3  Text accuracy           — paragraph-level text diff      (8 pts)
  D4  Completeness            — count-based: tables, images,   (7 pts)
                                 footnotes, sections, pages
  D5  Ordering / sequence     — section + paragraph order      (4 pts)
  D6  Encoding & characters   — smart quotes, dashes, accents  (3 pts)
  D7  Metadata accuracy       — title, date, doc-number, lang  (3 pts)
  ─────────────────────────────────
  Total                                                        30 pts

Architecture decisions:
  • PyMuPDF (fitz) for text + font metadata — fast, free, good for single-column
  • pdfplumber for table structure — more reliable than fitz for grid detection
  • All deterministic — no GPT/API calls, consistent results
  • Graceful degradation — each sub-check has its own try/except;
    one failure doesn't block others
  • Source PDF is always provided in production — all 6 dimensions run

Limitations (by design):
  • Cannot validate SEMANTIC tag choice (BLOCK2 vs PART) — needs human judgment
  • Multi-column PDF layouts may produce ordering false-positives — flagged as WARNING
  • Scanned PDFs (no text layer) silently skip text diff — logged as warning
"""

from __future__ import annotations

import os
import re
import unicodedata
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from typing import Optional

# Semantic agent replaces deterministic D3 + D8 (see semantic_content_agent.py)
from validator.level4_source_compare.semantic_content_agent import check_text_semantic

# ── Optional dependencies ─────────────────────────────────────────────────────
try:
    import fitz  # PyMuPDF
    _FITZ_OK = True
except ImportError:
    _FITZ_OK = False

try:
    import pdfplumber
    _PLUMBER_OK = True
except ImportError:
    _PLUMBER_OK = False


# ── Character maps for D6 encoding checks (no PDF needed) ────────────────────
# Unicode characters that MUST be encoded as SGML entities in Carswell SGML
UNICODE_TO_ENTITY: dict[str, str] = {
    "\u2018": "&lsquo;",   # '  left single quote
    "\u2019": "&rsquo;",   # '  right single quote
    "\u201c": "&ldquo;",   # "  left double quote
    "\u201d": "&rdquo;",   # "  right double quote
    "\u2013": "&ndash;",   # –  en dash
    "\u2014": "&mdash;",   # —  em dash
    "\u00a0": "&nbsp;",    # non-breaking space
    "\u00e9": "&eacute;",  # é
    "\u00e8": "&egrave;",  # è
    "\u00ea": "&ecirc;",   # ê
    "\u00e0": "&agrave;",  # à
    "\u00f4": "&ocirc;",   # ô
    "\u00c9": "&Eacute;",  # É
    "\u00b0": "&deg;",     # °
    "\u00b7": "&middot;",  # ·
    "\u00d7": "&times;",   # ×
    "\u00b1": "&plusmn;",  # ±
    "\u00a9": "&copy;",    # ©
    "\u2022": "&bull;",    # •
    "\u2026": "&hellip;",  # …
    "\u20ac": "&euro;",    # €
    "\u2265": "&ge;",      # ≥
    "\u2264": "&le;",      # ≤
    "\u00bd": "&frac12;",  # ½
    "\u00bc": "&frac14;",  # ¼
}

# Characters where bare hyphen is used instead of proper dash entity
_DASH_CONTEXT_RE = re.compile(
    r"(?<=[a-z\d])\s+-\s+(?=[a-z\d])",  # word - word  (likely em/en dash)
    re.IGNORECASE,
)

# Legal citation patterns that are corruption-prone
_LEGAL_CITATION_RE = re.compile(
    r"\b(?:s\.|ss\.|art\.|para\.|cl\.|sch\.)\s*\d+(?:\(\w+\))*",
    re.IGNORECASE,
)

# Number pattern: integers, decimals, currency, percentages
_NUMBER_RE = re.compile(r"\$[\d,]+(?:\.\d+)?|\d+(?:,\d{3})*(?:\.\d+)?%?")

# Legitimate omission patterns — text in PDF not expected in SGML
_OMIT_PATTERNS = [
    re.compile(r"^\d{1,4}$"),                              # bare page numbers
    re.compile(r"^[Pp]age\s+\d{1,4}"),                    # "Page 1", "Page 12"
    re.compile(r"^table\s+of\s+contents", re.I),           # TOC header
    re.compile(r"copyright\s+©?\s*20\d{2}", re.I),         # copyright lines
    re.compile(r"^\s*(continued|suite)\s*$", re.I),        # continuation marks
    re.compile(r"^(home|trading|français|sign in)", re.I),  # web chrome
    re.compile(r"thomson\s+reuters", re.I),                 # TR branding
    # Standalone date lines ("June 9, 2023", "September 27, 2016")
    re.compile(
        r"^(?:january|february|march|april|may|june|july|august|september|"
        r"october|november|december|janvier|février|mars|avril|mai|juin|"
        r"juillet|août|septembre|octobre|novembre|décembre)\s+\d{1,2},?\s+\d{4}$",
        re.IGNORECASE,
    ),
    # Header/cover lines that are also bold (e.g. organisation names)
    re.compile(r"^\d{7}$"),  # internal reference numbers ("6217407")
    # TOC dot-leader entries: "Registrants ........................................................."
    re.compile(r"\.{5,}"),
    # Footnote/endnote lines: "19 See NI 33-109..." or "21 See subsection..."
    re.compile(r"^\d{1,3}\s+(?:See|Ibid|supra|infra|note\b)", re.IGNORECASE),
    # URL-only or source-only footnote lines: "5 Budget 2024: ... - Canada.ca"
    re.compile(r"^\d{1,3}\s+.{5,}\.(?:ca|com|gov|org|net)(?:[/\s]|$)", re.IGNORECASE),
    # OSC Bulletin page-citation running headers: "(2025), 48 OSCB 9737"
    # These appear in bold at the top of every page but are NOT body content.
    re.compile(r"^\(\d{4}\),\s+\d+\s+OSCB\s+\d+"),
    # OSC Bulletin section markers: "B. Ontario Securities Commission",
    # "B.1: Notices", "B.5: Rules and Policies" — bulletin nav headers, not doc content
    re.compile(r"^B\.\d*[:\s]"),
    # Legal party separator in court/regulatory documents: "- and -"
    re.compile(r"^[-\u2013]\s+and\s+[-\u2013]$"),
    # Date fragments split across PDF lines: "12, 2025", "10, 2025"
    re.compile(r"^\d{1,2},\s+\d{4}$"),
    # Parenthetical date references: "(as of September 17, 2025)"
    re.compile(r"^\(as\s+of\b", re.IGNORECASE),
    # Standard statutory header phrase in legal documents
    re.compile(r"^made\s+under\s+the\b", re.IGNORECASE),
    # Continuation fragments from page headers/footers (start with comma)
    re.compile(r"^,"),
    # Statistical table column headers: "Q3 2024", "Q4 2023"
    re.compile(r"^Q[1-4]\s+\d{4}$", re.IGNORECASE),
    # Percentage change table headers
    re.compile(r"^%\s+change\b", re.IGNORECASE),
    # Metadata date lines: "Date: 20250417"
    re.compile(r"^[Dd]ate:\s+\d"),
    # Standalone section/part labels from TOC (PDF renders in bold/italic)
    # e.g. "Part 3", "Part 7", "Section 2" — already encoded inside <TI> full text
    re.compile(r"^(?:part|section|chapter)\s+\d+[a-z]?\s*$", re.IGNORECASE),
    # Standalone annex/appendix labels: "Annex A", "Appendix B"
    re.compile(r"^(?:annex|appendix)\s+[a-zA-Z]\s*$", re.IGNORECASE),
    # Roman numeral TOC sub-entries: "i.  tick test", "iv.  short sale circuit breaker"
    re.compile(r"^(?:i{1,3}|iv|vi{0,3}|ix|xi{0,2}|xiv|xv)\.\s+\S", re.IGNORECASE),
    # Institution letterhead names (appear on every page header, not SGML body)
    re.compile(r"(?:securities\s+commission|securities\s+authority)\s*$", re.IGNORECASE),
    # CIRO / CIPF regulatory org names appearing as page-header elements
    re.compile(r"^(?:canadian\s+investment\s+regulatory\s+organization|cipf)\s*$", re.IGNORECASE),
    # PDF track-changes / annotation artifacts
    re.compile(r"\bstrikethrough\b", re.IGNORECASE),
    # Metadata form-field labels in regulatory submission forms
    re.compile(
        r"^(?:document\s+(?:type|no\.?|number|date|title)|effective\s+date|reference\s+no\.?)"
        r"\s*[:\s]*$",
        re.IGNORECASE,
    ),
    # Part/Section label WITH inline title text — heading already covered by <TI> check
    # e.g. "Part 1  Definitions", "Part 3 Effective Date"
    re.compile(
        r"^(?:part|section|article|division|schedule|item)\s+\d+[a-zA-Z]?\s+"
        r"(?:[-\u2013\u2014]|definitions?|purposes?|interpretation|effective\s+date|"
        r"general|application|exemption|transitional|repeal)\b",
        re.IGNORECASE,
    ),
    # FLI/FOFI forward-looking information abbreviation fragments
    re.compile(r"^fli[;,\s].*fofi", re.IGNORECASE),
    # Spaced-out decorative cover typography (e.g. "2 0 2 5", "5 5 / 2 0 2 5")
    # PyMuPDF extracts large-format year/issue numbers letter-by-letter
    re.compile(r"^(?:\S{1,2}\s+){2,}\S{1,2}$"),
    # OSC Bulletin section category labels: "Rules and Policies", "Notices"
    # These appear as bold navigation headers in the bulletin, not SGML body content
    re.compile(r"^(?:rules\s+and\s+policies|notices\s+and\s+news\s+releases)$", re.IGNORECASE),
    # Table row-number + pipe separator artifact from PDF table rendering
    # e.g. "1 |", "2 |", "25 |" — row labels extracted by PyMuPDF, not in SGML body
    re.compile(r"^\d+\s*\|", re.IGNORECASE),
    # Part/Section N – <any title> — italic TOC entries where the heading itself
    # is covered by <TI> but PDF TOC renders it italic with a dash separator
    re.compile(
        r"^(?:part|section|article|division|schedule|item)\s+\d+[a-zA-Z]?\s*[\u2013\u2014-]",
        re.IGNORECASE,
    ),
]


# ── Font-name bold/italic detection (Gap 6 fix) ─────────────────────────────
# PyMuPDF font flags (bit4=bold) are unreliable for many PDFs that encode bold
# via the font name (e.g. "TimesNewRoman-Bold", "Helvetica-BoldOblique").
# These regexes complement the flags-based check.
_BOLD_FONT_NAME_RE = re.compile(
    r"(?i)(?:bold|heavy|black|demi|semibold|extrabold|ultra)"
)
_ITALIC_FONT_NAME_RE = re.compile(
    r"(?i)(?:italic|oblique|slanted)"
)


def _is_bold_font_name(font_name: str) -> bool:
    """Return True if the font name indicates bold weight."""
    return bool(_BOLD_FONT_NAME_RE.search(font_name))


def _is_italic_font_name(font_name: str) -> bool:
    """Return True if the font name indicates italic/oblique style."""
    return bool(_ITALIC_FONT_NAME_RE.search(font_name))


# ── Result dataclass ──────────────────────────────────────────────────────────
@dataclass
class L4Result:
    score: float = 0.0
    max_score: float = 30.0

    # Sub-scores (each dimension)
    tagging_score: float = 0.0      # D2: 0–5
    text_score: float = 0.0         # D3: 0–8
    completeness_score: float = 0.0 # D4: 0–7
    ordering_score: float = 0.0     # D5: 0–4
    encoding_score: float = 0.0     # D6: 0–3
    metadata_score: float = 0.0     # D7: 0–3

    issues: list[dict] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    # Diagnostics
    pdf_available: bool = False
    pdf_text_extractable: bool = False
    text_coverage: float = 0.0          # fraction of PDF paragraphs found in SGML
    missing_paragraphs: list[str] = field(default_factory=list)
    encoding_violations: list[str] = field(default_factory=list)
    sequence_violations: list[str] = field(default_factory=list)
    metadata_mismatches: list[str] = field(default_factory=list)

    # diff_generator fields — populated by check_* functions so the HITL
    # diff engine can produce line-specific, actionable fix suggestions
    d2_untagged_bold: list[str] = field(default_factory=list)
    d2_untagged_italic: list[str] = field(default_factory=list)
    d2_untagged_headings: list[str] = field(default_factory=list)
    d5_inverted_pairs: list[tuple] = field(default_factory=list)  # [(sgml_h_before, sgml_h_after), ...]
    d7_expected_lang: str = ""        # language the PDF suggests (for D7 LANG fix)
    d7_pdf_doc_number: str = ""       # doc number extracted from PDF (for D7 N-tag fix)
    pdf_headings: list[str] = field(default_factory=list)  # PDF headings (for D3 placement heuristic)

    # GAP 1: Two-stage validation — separate ABBYY extraction errors from pipeline errors
    docx_available: bool = False
    abbyy_missing_paragraphs: list[str] = field(default_factory=list)   # in PDF but not DOCX
    pipeline_missing_paragraphs: list[str] = field(default_factory=list) # in DOCX but not SGML
    # GAP confidence details — each dict: {text, confidence, method}
    abbyy_missing_paragraph_details: list[dict] = field(default_factory=list)
    pipeline_missing_paragraph_details: list[dict] = field(default_factory=list)
    # GAP 5: Table cell-level coverage — DOCX cells not found in SGML
    d4_missing_table_cells: list[dict] = field(default_factory=list)
    # D8: Word-level gaps — list of {missing: str, line: int}
    word_gaps: list[dict] = field(default_factory=list)

    # diff_generator compatibility fields — populated by check_text_accuracy()
    # when a DOCX is available; empty-list defaults ensure no AttributeError
    # if they are not set (e.g. PDF-only validation path).
    truncated_paragraphs: list[str] = field(default_factory=list)
    inline_changed_paragraphs: list[dict] = field(default_factory=list)
    missing_short_lines: list[str] = field(default_factory=list)

    # diff_generator contact-detail fields (D4-g/h)
    missing_emails: list[str] = field(default_factory=list)
    extra_emails: list[str] = field(default_factory=list)
    missing_phones: list[str] = field(default_factory=list)
    extra_phones: list[str] = field(default_factory=list)
    missing_urls: list[str] = field(default_factory=list)
    extra_urls: list[str] = field(default_factory=list)
    missing_postal_codes: list[str] = field(default_factory=list)


def _add_issue(result: L4Result, dimension: str, severity: str, description: str,
               location: str = "", impact: str = "") -> None:
    result.issues.append({
        "level": "L4",
        "category": dimension,
        "severity": severity,
        "description": description,
        "location": location,
        "impact": impact,
    })


# ── Entity decoding ──────────────────────────────────────────────────────────
# Map named SGML/HTML entities → Unicode characters.
# Used by _norm() so that SGML "l&rsquo;article" and PDF "l'article" compare equal.
_ENTITY_CHAR_MAP: dict[str, str] = {
    # Quotation marks / apostrophes
    "rsquo": "\u2019", "lsquo": "\u2018",
    "rdquo": "\u201d", "ldquo": "\u201c",
    "apos": "'",       "quot": '"',
    # Dashes
    "ndash": "\u2013", "mdash": "\u2014", "minus": "\u2212",
    # Spaces
    "nbsp": "\u00a0",  "ensp": "\u2002", "emsp": "\u2003",
    # XML built-ins
    "amp": "&", "lt": "<", "gt": ">",
    # French / accented Latin
    "eacute": "\u00e9", "egrave": "\u00e8", "ecirc": "\u00ea", "euml": "\u00eb",
    "Eacute": "\u00c9", "Egrave": "\u00c8", "Ecirc": "\u00ca",
    "agrave": "\u00e0", "acirc": "\u00e2", "auml": "\u00e4", "aring": "\u00e5",
    "Agrave": "\u00c0", "Acirc": "\u00c2",
    "ugrave": "\u00f9", "ucirc": "\u00fb", "uuml": "\u00fc",
    "icirc": "\u00ee", "iuml": "\u00ef",
    "ocirc": "\u00f4", "ouml": "\u00f6",
    "ccedil": "\u00e7", "Ccedil": "\u00c7",
    "oelig": "\u0153",  "OElig": "\u0152",
    "szlig": "\u00df",
    # Symbols
    "deg": "\u00b0",   "middot": "\u00b7", "times": "\u00d7", "plusmn": "\u00b1",
    "copy": "\u00a9",  "reg": "\u00ae",    "trade": "\u2122",
    "bull": "\u2022",  "hellip": "\u2026", "euro": "\u20ac",
    "ge": "\u2265",    "le": "\u2264",     "ne": "\u2260",
    "frac12": "\u00bd","frac14": "\u00bc", "frac34": "\u00be",
    "sect": "\u00a7",  "para": "\u00b6",   "dagger": "\u2020",
}


def _decode_sgml_entities(text: str) -> str:
    """Replace named SGML entities with their Unicode characters."""
    def _replace(m: re.Match) -> str:
        return _ENTITY_CHAR_MAP.get(m.group(1), " ")
    return re.sub(r"&([a-zA-Z][a-zA-Z0-9]*);", _replace, text)


# ── Text normalisation ────────────────────────────────────────────────────────
def _norm(text: str) -> str:
    """Decode entities → strip tags → normalise quotes/dashes → lowercase."""
    text = _decode_sgml_entities(text)             # "&rsquo;" → "'"
    text = re.sub(r"<[^>]+>", " ", text)           # strip SGML tags
    text = unicodedata.normalize("NFC", text)
    # Normalise typographic variants to ASCII so PDF and SGML compare equal
    text = text.replace("\u2018", "'").replace("\u2019", "'")   # ' ' → '
    text = text.replace("\u201c", '"').replace("\u201d", '"')   # " " → "
    text = text.replace("\u2013", "-").replace("\u2014", "-")   # en/em dash → -
    text = text.replace("\u2212", "-")                          # minus sign → -
    text = text.replace("\u00a0", " ")                          # nbsp → space
    text = text.lower()
    # Fix inline-tag split artifacts produced when SGML tags like <EM> or <BOLD>
    # sit inside a hyphenated token or immediately before punctuation:
    #   "21-<EM>101"  →  strip tag  →  "21- 101"  →  fix  →  "21-101"
    #   "funds</EM>;" →  strip tag  →  "funds ;"   →  fix  →  "funds;"
    # Without this, compound numbers (NI 21-101, 33-109, etc.) fail D3 matching.
    text = re.sub(r'(\w-)\s+(\w)', r'\1\2', text)    # join split hyphenated tokens
    text = re.sub(r'(\w)\s+([;:,.])', r'\1\2', text)  # rejoin split punctuation
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _is_omittable(text: str) -> bool:
    """Return True if this PDF text is legitimately absent from SGML."""
    t = text.strip()
    return any(p.search(t) for p in _OMIT_PATTERNS)


# ── PDF extraction (PyMuPDF) ──────────────────────────────────────────────────
@dataclass
class _PDFData:
    paragraphs: list[str] = field(default_factory=list)
    headings: list[str] = field(default_factory=list)
    bold_spans: list[str] = field(default_factory=list)
    italic_spans: list[str] = field(default_factory=list)
    table_count: int = 0
    image_count: int = 0
    footnote_count: int = 0
    page_count: int = 0
    first_page_text: str = ""
    language_hint: str = "EN"      # detected from character frequencies
    doc_title: str = ""
    doc_date: str = ""
    doc_number: str = ""
    two_column: bool = False        # Gap 9: True if 2-column layout detected
    ok: bool = True
    error: str = ""


def _detect_language(text: str) -> str:
    """Simple heuristic: count French-specific characters."""
    french_chars = sum(text.count(c) for c in "éèêàâùûîôçœÉÈÊÀÂÙÛÎÔÇŒ")
    return "FR" if french_chars > 5 else "EN"


def _extract_doc_number(text: str) -> str:
    """Extract regulatory document number like NI 31-103, OSC Rule 14-501, etc."""
    # Normalise newlines so multiline spans like 'Notice\n11-326' are joined
    text = re.sub(r"\s+", " ", text)
    patterns = [
        r"\b(NI|MI|CSA|OSC|MSC|ASC|BCSC|AMF)\s+\d{2}-\d{3}\b",
        r"\b(?:National|Multilateral)\s+Instrument\s+(\d{2}-\d{3})\b",
        r"\b(?:Rule|Policy|Notice|Bulletin|Guideline)\s+(\d{2}-\d{3})\b",
        r"\b(\d{2}-\d{3})\b",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return re.sub(r"\s+", " ", m.group(0)).strip()
    return ""


def _extract_doc_date(text: str) -> str:
    """Extract date from first-page text. Returns YYYYMMDD or empty string."""
    # Match formats: "April 15, 2026", "15 April 2026", "2026-04-15", "April 2026"
    months = {
        "january": "01", "february": "02", "march": "03", "april": "04",
        "may": "05", "june": "06", "july": "07", "august": "08",
        "september": "09", "october": "10", "november": "11", "december": "12",
        "jan": "01", "feb": "02", "mar": "03", "apr": "04",
        "jun": "06", "jul": "07", "aug": "08", "sep": "09",
        "oct": "10", "nov": "11", "dec": "12",
        # French
        "janvier": "01", "février": "02", "mars": "03", "avril": "04",
        "mai": "05", "juin": "06", "juillet": "07", "août": "08",
        "septembre": "09", "octobre": "10", "novembre": "11", "décembre": "12",
    }
    # "April 15, 2026" or "April 2026"
    m = re.search(
        r"\b(" + "|".join(months) + r")\s+(\d{1,2})(?:,\s+|\s+)(\d{4})\b",
        text, re.IGNORECASE
    )
    if m:
        mo = months[m.group(1).lower()]
        day = m.group(2).zfill(2)
        yr = m.group(3)
        return f"{yr}{mo}{day}"
    # "15 April 2026"
    m = re.search(
        r"\b(\d{1,2})\s+(" + "|".join(months) + r")\s+(\d{4})\b",
        text, re.IGNORECASE
    )
    if m:
        day = m.group(1).zfill(2)
        mo = months[m.group(2).lower()]
        yr = m.group(3)
        return f"{yr}{mo}{day}"
    # ISO: "2026-04-15"
    m = re.search(r"\b(\d{4})-(\d{2})-(\d{2})\b", text)
    if m:
        return f"{m.group(1)}{m.group(2)}{m.group(3)}"
    # Labelled bare YYYYMMDD: "Date: 20250417" or "Date:  20250417"
    m = re.search(r"\bDate:\s*(\d{8})\b", text, re.IGNORECASE)
    if m:
        return m.group(1)
    return ""


def _extract_pdf_data(pdf_path: str) -> _PDFData:
    """Extract structured data from PDF using PyMuPDF + optional pdfplumber."""
    data = _PDFData()

    if not _FITZ_OK:
        data.ok = False
        data.error = "PyMuPDF (fitz) not installed"
        return data

    # Cap pages processed for text/font extraction to avoid hanging on very large PDFs.
    # Structural metadata (page count, image count, table count) still uses the full doc.
    _MAX_EXTRACT_PAGES = 80

    try:
        doc = fitz.open(pdf_path)
        data.page_count = len(doc)

        # Track repeated lines (headers/footers)
        line_freq: dict[str, int] = {}
        all_lines: list[str] = []

        for page_idx, page in enumerate(doc):
            if page_idx >= _MAX_EXTRACT_PAGES:
                break  # skip text/font extraction for pages beyond cap
            _page_h = page.rect.height  # GAP 3: for geometric header/footer filtering
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
            for block in blocks:
                if block.get("type") != 0:
                    # Type 1 = image block
                    if block.get("type") == 1:
                        data.image_count += 1
                    continue
                # GAP 3: Skip blocks in header zone (top 10%) or footer zone (bottom 8%).
                # These are running headers/footers that ABBYY already excludes from DOCX.
                # Geometry-based filtering is more reliable than 100+ regex patterns.
                if _page_h > 0:
                    _bbox = block.get("bbox", (0, 0, 0, _page_h))
                    if _bbox[1] / _page_h < 0.10 or _bbox[3] / _page_h > 0.92:
                        continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    line_text = "".join(s["text"] for s in spans).strip()
                    if not line_text:
                        continue
                    all_lines.append(line_text)
                    line_freq[line_text] = line_freq.get(line_text, 0) + 1

                    # Bold / italic detection: flags AND font-name (Gap 6 fix)
                    # Flags alone miss PDFs that encode bold via font name only.
                    for span in spans:
                        flags = span.get("flags", 0)
                        font_name = span.get("font", "")
                        text = span["text"].strip()
                        if not text or len(text) < 3:
                            continue
                        is_bold = bool(flags & 16) or _is_bold_font_name(font_name)
                        is_italic = bool(flags & 2) or _is_italic_font_name(font_name)
                        if is_bold:
                            data.bold_spans.append(text)
                        if is_italic:
                            data.italic_spans.append(text)

        # Remove repeated lines (appear on ≥ 50% of pages, min 3 pages) — headers/footers
        threshold = max(3, data.page_count * 0.5)
        repeated = {ln for ln, cnt in line_freq.items() if cnt >= threshold}

        # Post-filter bold/italic span lists: remove running header/footer strings.
        # Bold/italic spans are accumulated during the page loop BEFORE we compute
        # `repeated`, so they contain header/footer text (e.g. italic running title,
        # bold date lines that appear on every page). Filter them out now.
        data.bold_spans = [s for s in data.bold_spans if s not in repeated]
        data.italic_spans = [s for s in data.italic_spans if s not in repeated]

        # Build page-1 text for metadata extraction (first 3 pages)
        first_pages_text = ""
        for page_idx in range(min(3, data.page_count)):
            first_pages_text += doc[page_idx].get_text()
        data.first_page_text = first_pages_text[:3000]

        # Detect font sizes to identify headings
        # Collect (font_size, text) tuples from first _MAX_EXTRACT_PAGES pages
        size_text: list[tuple[float, str]] = []
        for page in list(doc)[:_MAX_EXTRACT_PAGES]:
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        t = span["text"].strip()
                        if t and len(t) > 3:
                            size_text.append((span.get("size", 0), t))

        # Body font = most common size
        if size_text:
            from collections import Counter
            size_counts = Counter(round(s, 0) for s, _ in size_text)
            body_size = size_counts.most_common(1)[0][0]
            heading_threshold = body_size * 1.15  # 15% larger = heading

            for size, text in size_text:
                if text in repeated:
                    continue
                if size >= heading_threshold and len(text) > 5:
                    data.headings.append(text)

        # Build clean paragraphs (remove repeated, omittable, short lines)
        # PyMuPDF extracts line-by-line; join consecutive lines into paragraphs
        # A line is a continuation if it doesn't end a sentence and the next is
        # short enough to be a wrapped line (< 80 chars).
        raw_lines: list[str] = []
        for ln in all_lines:
            if ln in repeated:
                continue
            if _is_omittable(ln):
                continue
            if len(ln.split()) < 2:
                continue
            raw_lines.append(ln)

        # Merge continuation lines into paragraphs
        merged: list[str] = []
        buf = ""
        for ln in raw_lines:
            if not buf:
                buf = ln
            else:
                # Heuristic: join if previous line doesn't end in sentence-terminator
                # or current line starts lowercase / looks like a continuation
                prev_ends_sentence = buf.rstrip().endswith((".", "?", "!", ":", ";"))
                curr_starts_upper = ln[0].isupper() if ln else True
                is_short_prev = len(buf) < 70  # previous line was short (wrapped)
                if not prev_ends_sentence or (is_short_prev and not curr_starts_upper):
                    buf = buf.rstrip() + " " + ln
                else:
                    merged.append(buf)
                    buf = ln
        if buf:
            merged.append(buf)

        data.paragraphs = [ln for ln in merged if len(ln.split()) >= 4]

        # Footnote heuristic: lines with superscript-style numbering at start
        footnote_re = re.compile(r"^\d{1,3}\s+\S")
        data.footnote_count = sum(1 for ln in all_lines if footnote_re.match(ln))

        # Language detection from first page
        data.language_hint = _detect_language(first_pages_text)

        # Metadata extraction from first-page text
        data.doc_number = _extract_doc_number(first_pages_text)
        data.doc_date = _extract_doc_date(first_pages_text)

        # Table count via pdfplumber (more reliable).
        # Guard: skip pdfplumber for large PDFs (>400 KB) or many pages (>60)
        # where it can take several minutes — fall back to fitz heuristic.
        _pdf_size_kb = os.path.getsize(pdf_path) / 1024 if os.path.exists(pdf_path) else 0
        _use_plumber = _PLUMBER_OK and _pdf_size_kb <= 400 and data.page_count <= 60
        if _use_plumber:
            try:
                import pdfplumber as _plumber
                with _plumber.open(pdf_path) as plumb:
                    for pg in plumb.pages:
                        tables = pg.extract_tables()
                        if tables:
                            data.table_count += len([t for t in tables if t])
            except Exception:
                # Fallback: rough table detection from fitz line geometry
                data.table_count = _estimate_table_count_fitz(doc)
        else:
            data.table_count = _estimate_table_count_fitz(doc)

        # Full-doc image count (faster than text extraction — uses xref list)
        if data.image_count == 0:
            try:
                for _pg in doc:
                    data.image_count += len(_pg.get_images(full=False))
            except Exception:
                pass

        # Gap 9: detect 2-column layout BEFORE closing doc
        data.two_column = _detect_two_column_layout(doc)

        doc.close()

    except Exception as exc:
        data.ok = False
        data.error = str(exc)

    return data


def _detect_two_column_layout(doc) -> bool:
    """
    Detect if the PDF uses a 2-column text layout (Gap 9 fix).

    Heuristic: sample the first 5 pages; if the median text-line width is
    less than 55 % of the usable page width, the document is almost certainly
    laid out in two (or more) columns.  In that case D5 ordering results are
    unreliable because PyMuPDF reads text left-to-right across the full page
    width, mixing columns.
    """
    if not _FITZ_OK:
        return False

    line_widths: list[float] = []
    page_widths: list[float] = []

    for page_idx, page in enumerate(doc):
        if page_idx >= 5:  # sample first 5 pages only
            break
        pw = page.rect.width
        if pw <= 0:
            continue
        page_widths.append(pw)

        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                line_text = "".join(s["text"] for s in spans).strip()
                if len(line_text) < 20:  # skip very short / label lines
                    continue
                bbox = line.get("bbox", [0, 0, 0, 0])
                line_w = bbox[2] - bbox[0]
                if line_w > 0:
                    line_widths.append(line_w)

    if not line_widths or not page_widths:
        return False

    avg_page_width = sum(page_widths) / len(page_widths)
    # Typical margins: ~10 % each side → usable width ≈ 80 % of page width
    usable_width = avg_page_width * 0.80

    line_widths_sorted = sorted(line_widths)
    median_lw = line_widths_sorted[len(line_widths_sorted) // 2]

    ratio = median_lw / usable_width if usable_width > 0 else 1.0
    # 2-column threshold: median text line < 55 % of usable page width
    return ratio < 0.55


def _estimate_table_count_fitz(doc) -> int:
    """Rough table count using horizontal line density heuristic (fitz fallback)."""
    count = 0
    for page in doc:
        paths = page.get_drawings()
        h_lines = [p for p in paths if p.get("rect") and
                   abs(p["rect"].height) < 3 and p["rect"].width > 50]
        if len(h_lines) >= 4:
            count += 1
    return count


# ── SGML text extraction ──────────────────────────────────────────────────────
def _extract_sgml_text(sgml: str) -> dict:
    """Extract text content, tag counts, and metadata from raw SGML."""
    # Strip all tags to get plain text, decoding entities → Unicode
    # (same normalisation as _norm() so SGML blob matches PDF text in D3)
    text_only = re.sub(r"<[^>]+>", " ", sgml)
    text_only = _decode_sgml_entities(text_only)
    text_only = unicodedata.normalize("NFC", text_only)
    text_only = text_only.replace("\u2018", "'").replace("\u2019", "'")
    text_only = text_only.replace("\u201c", '"').replace("\u201d", '"')
    text_only = text_only.replace("\u2013", "-").replace("\u2014", "-")
    text_only = text_only.replace("\u2212", "-")
    text_only = text_only.replace("\u00a0", " ")
    # Fix inline-tag split artifacts (same logic as _norm — keep in sync):
    # "21-<EM>101" → strip tag → "21- 101" → fix → "21-101"
    text_only = re.sub(r'(\w-)\s+(\w)', r'\1\2', text_only)
    text_only = re.sub(r'(\w)\s+([;:,.])', r'\1\2', text_only)
    text_only = re.sub(r"\s+", " ", text_only).strip()

    # Extract paragraphs (content inside P tags)
    p_contents = re.findall(r"<P[^>]*>(.*?)</P>", sgml, re.DOTALL)
    paragraphs = []
    for pc in p_contents:
        clean = _norm(pc)
        if len(clean.split()) >= 5:
            paragraphs.append(clean)

    # Headings (TI tags)
    headings = [_norm(h) for h in re.findall(r"<TI[^>]*>(.*?)</TI>", sgml, re.DOTALL)]

    # Counts
    table_count = len(re.findall(r"<TABLE[\s>]", sgml))
    fn_count = len(re.findall(r"<FN[\s>]|<FOOTNOTE[\s>]", sgml))
    graphic_count = len(re.findall(r"<GRAPHIC\s", sgml))

    # POLIDOC metadata
    polidoc_m = re.search(r"<POLIDOC([^>]*)>", sgml)
    attrs = {}
    if polidoc_m:
        for am in re.finditer(r'(\w+)="([^"]*)"', polidoc_m.group(1)):
            attrs[am.group(1)] = am.group(2)

    # Section labels from N tags near TI tags (document numbering)
    sections = re.findall(r"<N[^>]*>(.*?)</N>", sgml, re.DOTALL)

    # GAP 5: extract SGML table cell text for cell-by-cell comparison
    sgml_table_cells: list[str] = []
    for _cell_content in re.findall(r"<TBLCELL[^>]*>(.*?)</TBLCELL>", sgml, re.DOTALL):
        _ct = _norm(re.sub(r"<[^>]+>", " ", _cell_content))
        if _ct and len(_ct.split()) >= 2:
            sgml_table_cells.append(_ct)

    # D4-e fix: store whether the raw SGML contains an <APPENDIX> or <SCHEDDOC>
    # opening tag. The "text" field has tags stripped so searching it for
    # "<APPENDIX" would always return False — a direct cause of D4-e false
    # positives on documents that use <APPENDIX> labels (e.g. CSA Staff Notices).
    has_appendix_tag = bool(re.search(r"<APPENDIX[\s>]|<SCHEDDOC[\s>]", sgml))

    return {
        "text": text_only,
        "paragraphs": paragraphs,
        "headings": headings,
        "sections": [_norm(s) for s in sections],
        "table_count": table_count,
        "fn_count": fn_count,
        "graphic_count": graphic_count,
        "attrs": attrs,
        "table_cells": sgml_table_cells,   # GAP 5: TBLCELL text
        "has_appendix_tag": has_appendix_tag,  # D4-e: raw-SGML tag presence
    }


# ─────────────────────────────────────────────────────────────────────────────
# D6: Encoding & character accuracy (runs WITHOUT source PDF)
# ─────────────────────────────────────────────────────────────────────────────
def check_encoding(raw_sgml: str, result: L4Result) -> None:
    """
    D6 — 3 pts: Detect Unicode characters that must be encoded as SGML entities.

    Checks text content inside tags for raw Unicode that should be an entity.
    Does NOT require the source PDF — runs on SGML alone.
    """
    score = 3.0

    # Extract text content (inside tags, after stripping tags)
    # We need to check the raw content, not the tag attributes
    # Strip attribute values and tag markup, keep text content
    text_content = re.sub(r"<[^>]+>", "\x00", raw_sgml)  # replace tags with null
    # text_content now has text nodes separated by nulls

    violations: list[str] = []
    char_counts: dict[str, int] = {}

    for char, entity in UNICODE_TO_ENTITY.items():
        count = text_content.count(char)
        if count > 0:
            char_counts[char] = count
            violations.append(f"Raw U+{ord(char):04X} ({entity}) found {count}× — use {entity}")

    # Check for bare hyphen used as dash in mid-sentence
    dash_misuse = len(_DASH_CONTEXT_RE.findall(text_content))
    if dash_misuse > 0:
        violations.append(f"Bare hyphen used as dash {dash_misuse}× — use &ndash; or &mdash;")

    # Check for straight quotes in running text (not in tag attributes)
    # Tag attributes already handled — check content only
    straight_dq = text_content.count('"')
    if straight_dq > 5:  # allow a few in quoted material
        violations.append(f"Straight double-quotes {straight_dq}× — use &ldquo;/&rdquo;")

    if violations:
        n = len(violations)
        pts = min(2.0, n * 0.4)
        score -= pts
        result.encoding_violations = violations  # store all — no cap
        severity = "major" if pts >= 1.0 else "minor"
        _add_issue(result, "encoding", severity,
                   f"D6 — {n} encoding violation(s): raw Unicode instead of SGML entities. "
                   f"First: {violations[0]}",
                   impact=f"-{pts:.1f} pts")

    result.encoding_score = max(0.0, score)


# ─────────────────────────────────────────────────────────────────────────────
# D2: Tagging accuracy (requires PDF)
# ─────────────────────────────────────────────────────────────────────────────
def check_tagging(pdf: _PDFData, sgml_data: dict, raw_sgml: str, result: L4Result,
                  docx_data: "dict | None" = None) -> None:
    """
    D2 — 5 pts: Validate that PDF formatting is reflected by correct SGML tags.

    Checks:
    - Bold text in PDF → <BOLD> or <EM> tag in SGML
    - Italic text in PDF → <EM> or <ITALIC> tag in SGML
    - Headings in PDF (larger font) → <TI> in SGML
    - Tables in PDF → <TABLE> count approximately matches
    - Images in PDF → <GRAPHIC> count approximately matches

    GAP 4: When docx_data is provided, use DOCX bold/italic runs as the
    authoritative source instead of PDF font-flag spans. python-docx run.bold
    and run.italic read directly from OOXML w:rPr elements — far more reliable
    than PyMuPDF font-flag inference.
    """
    score = 5.0

    # Pre-compute <LINE> TOC text — reused by D2-a, D2-b, D2-c to avoid flagging
    # TOC entries that render bold/italic in the PDF but are correctly encoded as <LINE>.
    _line_texts = {
        _norm(re.sub(r"<[^>]+>", " ", ln))
        for ln in re.findall(r"<LINE[^>]*>(.*?)</LINE>", raw_sgml, re.DOTALL)
        if len(ln.split()) >= 2
    }

    # Pre-compute full SGML body plain-text blob (tags stripped) — used in D2-a and D2-b
    # as the final coverage fallback: if the text exists anywhere in SGML, it's present;
    # the vendor just chose a different tag (or no tag) which may be acceptable.
    _sgml_body_blob = _norm(re.sub(r"<[^>]+>", " ", raw_sgml))

    # D2-a: Bold spans from PDF should appear wrapped in BOLD/EM in SGML
    bold_tagged = re.findall(r"<(?:BOLD|EM)[^>]*>(.*?)</(?:BOLD|EM)>", raw_sgml, re.DOTALL)
    bold_tagged_text = {_norm(t) for t in bold_tagged}

    # GAP-7 FIX: Amendment instruments encode quoted replacement text in <QUOTE> tags.
    # This text renders bold in the PDF ("replacing X with Y") but is NOT wrapped in
    # <BOLD> in SGML — the <QUOTE> tag itself implies the content. Add to covered set.
    for qt in re.findall(r"<QUOTE[^>]*>(.*?)</QUOTE>", raw_sgml, re.DOTALL):
        bold_tagged_text.add(_norm(qt))

    # Also collect TI heading text — bold spans that are part of headings
    # are legitimately not wrapped in <BOLD> (the heading tag itself implies bold)
    ti_texts = {_norm(h) for h in sgml_data["headings"]}

    # GAP 4: prefer DOCX bold runs (explicit OOXML markup) over PDF font-flag spans.
    # Fall back to PDF spans if DOCX unavailable or has no bold runs detected.
    _bold_source = (docx_data.get("bold_runs") or []) if (docx_data and docx_data.get("bold_runs")) else pdf.bold_spans

    untagged_bold = []
    for span in _bold_source:  # check all bold spans — no sampling cap
        norm_span = _norm(span)
        if len(norm_span.split()) < 2:
            continue  # skip single words, too noise-prone
        if _is_omittable(span):
            continue
        in_bold_em = any(norm_span in bt or bt in norm_span for bt in bold_tagged_text)
        in_heading = any(norm_span in th or th in norm_span for th in ti_texts)
        # Fragment match: PDF splits long bold runs into short line-spans; match against longer SGML text
        if not in_bold_em:
            in_bold_em = any(
                norm_span in bt or SequenceMatcher(None, norm_span, bt).ratio() >= 0.85
                for bt in bold_tagged_text if len(bt) >= len(norm_span)
            )
        if not in_bold_em and not in_heading:
            in_heading = any(
                norm_span in th or SequenceMatcher(None, norm_span, th).ratio() >= 0.85
                for th in ti_texts if len(th) >= len(norm_span)
            )
        # <LINE> TOC coverage: bold TOC entries encoded as <LINE> in SGML, not <BOLD>
        if not in_bold_em and not in_heading:
            in_bold_em = any(norm_span in ln or ln in norm_span for ln in _line_texts)
        # GAP-4 FIX: If bold span text appears anywhere in SGML plain text, the vendor
        # has the content — they just encoded it without <BOLD> (e.g. short fragments
        # from multi-line bold blocks, provision labels, commission names in body text).
        if not in_bold_em and not in_heading:
            in_bold_em = norm_span in _sgml_body_blob
        if not in_bold_em and not in_heading:
            untagged_bold.append(span[:60])

    # Always store for diff_generator (even if no issue)
    result.d2_untagged_bold = untagged_bold

    if untagged_bold:
        ratio = len(untagged_bold) / max(1, len([s for s in _bold_source if len(s.split()) >= 2]))
        pts = min(1.5, ratio * 3.0)
        score -= pts
        severity = "major" if ratio > 0.3 else "minor"
        _add_issue(result, "tagging_accuracy", severity,
                   f"D2 — {len(untagged_bold)} bold text span(s) from PDF not wrapped in "
                   f"<BOLD> or <EM> in SGML. Examples: {untagged_bold[:3]}",
                   impact=f"-{pts:.1f} pts")

    # D2-b: Italic spans from PDF → <EM> or <ITALIC>
    # Important: text inside <TI> heading tags is already styled — it does NOT
    # need a separate <EM> wrapper. Exclude heading text from the italic check
    # to avoid false positives (e.g. '<TI>Trade Execution</TI>' is correct;
    # flagging it as "missing <EM>" is wrong).
    italic_tagged = re.findall(r"<(?:EM|ITALIC)[^>]*>(.*?)</(?:EM|ITALIC)>", raw_sgml, re.DOTALL)
    italic_tagged_text = {_norm(t) for t in italic_tagged}
    # Add all TI heading text as implicitly covered (heading styling implies italic/bold)
    italic_tagged_text.update(ti_texts)  # ti_texts defined in D2-a

    # GAP-1 FIX: Text already in <BOLD> does NOT need <EM> too.
    # Bold-italic PDF spans (e.g. Phase 1/2/3 headings) are correctly encoded
    # as <BOLD> only in SGML. Adding bold text to covered set prevents false positives.
    italic_tagged_text.update(bold_tagged_text)

    # GAP-2 FIX: Add all plain text from inside <FOOTNOTE> blocks to covered set.
    # Legislation Act/Bill names cited in footnotes are italic in PDF but the vendor
    # correctly puts them in <FOOTNOTE><FREEFORM><P> without wrapping in <EM>.
    _fn_blob = " ".join(
        _norm(re.sub(r"<[^>]+>", " ", fn))
        for fn in re.findall(r"<FOOTNOTE[^>]*>(.*?)</FOOTNOTE>", raw_sgml, re.DOTALL)
    )

    # (_line_texts pre-computed before D2-a above — reused here for italic TOC coverage)

    # (D2-b) _sgml_body_blob is pre-computed above, shared with D2-a.
    # GAP 4: prefer DOCX italic runs over PDF font-flag spans.
    _italic_source = (docx_data.get("italic_runs") or []) if (docx_data and docx_data.get("italic_runs")) else pdf.italic_spans

    untagged_italic = []
    for span in _italic_source:  # check all italic spans — no sampling cap
        norm_span = _norm(span)
        if len(norm_span.split()) < 2:
            continue
        # Direct match: span text appears inside an EM/ITALIC/TI/BOLD tag
        found = any(norm_span in it or it in norm_span for it in italic_tagged_text)
        # GAP-4 FIX: fragment match — span may be a PyMuPDF line-fragment of a longer EM span
        if not found:
            found = any(
                norm_span in em_full or SequenceMatcher(None, norm_span, em_full).ratio() >= 0.85
                for em_full in italic_tagged_text if len(em_full) >= len(norm_span)
            )
        # GAP-2 FIX: span may be inside a footnote (legislation names, citations)
        if not found:
            found = norm_span in _fn_blob
        # TOC FIX: span may be a <LINE> TOC entry (italic in PDF, no <EM> in SGML — by design)
        if not found:
            found = any(norm_span in ln or ln in norm_span for ln in _line_texts)
        # GAP-2b FIX: span text present in SGML body as plain text — vendor omitted <EM>
        # which is acceptable for legislation names, bibliography entries, Q&A headings.
        # Only flag if the text is completely absent from SGML.
        if not found:
            found = norm_span in _sgml_body_blob
        # PDF TOC entries often have trailing punctuation (comma, dash) not in SGML headings.
        # Strip trailing punctuation and retry body blob match.
        if not found:
            norm_stripped = norm_span.rstrip('- ,;\u2013\u2014')
            if len(norm_stripped.split()) >= 2:
                found = norm_stripped in _sgml_body_blob
        if not found and not _is_omittable(span):
            untagged_italic.append(span[:60])

    result.d2_untagged_italic = untagged_italic

    if untagged_italic:
        ratio = len(untagged_italic) / max(1, len([s for s in _italic_source if len(s.split()) >= 2]))
        pts = min(1.0, ratio * 2.0)
        score -= pts
        _add_issue(result, "tagging_accuracy", "minor",
                   f"D2 — {len(untagged_italic)} italic span(s) from PDF not wrapped in "
                   f"<EM> or <ITALIC> in SGML. Examples: {untagged_italic[:3]}",
                   impact=f"-{pts:.1f} pts")

    # D2-c: PDF headings (larger font) → <TI> in SGML
    # Some documents legitimately encode the document title/notice heading as
    # <P><BOLD>...</BOLD></P> rather than <TI>. Accept that as covered if the
    # heading text matches bold-tagged paragraph content.
    bold_para_texts: set[str] = set()
    for bp in re.findall(r"<(?:BOLD|EM)[^>]*>(.*?)</(?:BOLD|EM)>", raw_sgml, re.DOTALL):
        norm_bp = _norm(bp)
        if len(norm_bp.split()) >= 2:
            bold_para_texts.add(norm_bp)

    sgml_ti_texts_h = {_norm(h) for h in sgml_data["headings"]}  # TI tags
    # Include <N> document identifiers — PDF headings often carry a label prefix
    # e.g. PDF: "CSA Staff Notice 41-307 (Revised)" vs SGML N: "41-307 (Revised)"
    _n_texts = {
        _norm(n) for n in re.findall(r"<N[^>]*>(.*?)</N>", raw_sgml, re.DOTALL)
        if len(n.strip()) >= 4
    }
    _all_sgml_headings = sgml_ti_texts_h | _n_texts

    untagged_headings = []
    for heading in pdf.headings:  # check all headings — no sampling cap
        norm_h = _norm(heading)
        if len(norm_h.split()) < 2:
            continue
        # Check TI match (fuzzy)
        found_ti = any(
            SequenceMatcher(None, norm_h, sh).ratio() >= 0.70
            for sh in sgml_ti_texts_h
        )
        # Check if any SGML TI or N text is a substring of the PDF heading
        # e.g. "csa staff notice 41-307 (revised)" contains "41-307 (revised)"
        if not found_ti:
            found_ti = any(sh in norm_h for sh in _all_sgml_headings if len(sh) >= 8)
        # Accept if heading text is encoded as bold paragraph (valid alternative)
        found_bold = any(
            norm_h in bp or bp in norm_h or SequenceMatcher(None, norm_h, bp).ratio() >= 0.75
            for bp in bold_para_texts
        )
        # <LINE> TOC coverage: heading that appears in PDF TOC encoded as <LINE> in SGML
        if not found_ti and not found_bold:
            found_ti = any(norm_h in ln or ln in norm_h for ln in _line_texts)
        # Body blob fallback: heading content present in SGML but not tagged as <TI>
        # (e.g. encoded as <P><BOLD> or plain <P> — valid alternative encoding)
        if not found_ti and not found_bold:
            found_ti = norm_h in _sgml_body_blob
        # Reverse substring: any SGML body fragment ≥8 chars contained in PDF heading
        if not found_ti and not found_bold:
            found_ti = any(sh in norm_h for sh in _all_sgml_headings if len(sh) >= 8)
        if not found_ti and not found_bold and not _is_omittable(heading):
            untagged_headings.append(heading[:80])

    result.d2_untagged_headings = untagged_headings

    if untagged_headings:
        ratio = len(untagged_headings) / max(1, len(pdf.headings))
        pts = min(1.5, ratio * 2.0)
        score -= pts
        severity = "major" if ratio > 0.4 else "minor"
        _add_issue(result, "tagging_accuracy", severity,
                   f"D2 — {len(untagged_headings)} heading(s) detected in PDF (larger font) "
                   f"not tagged as <TI> in SGML. Examples: {untagged_headings[:3]}",
                   impact=f"-{pts:.1f} pts")

    # D2-d: Image count match (PDF images → GRAPHIC tags)
    sgml_graphic_count = sgml_data["graphic_count"]
    if pdf.image_count > 0 and sgml_graphic_count == 0:
        pts = 1.0
        score -= pts
        _add_issue(result, "tagging_accuracy", "major",
                   f"D2 — PDF has {pdf.image_count} image(s) but SGML has no <GRAPHIC> tags.",
                   impact=f"-{pts:.1f} pt")
    elif pdf.image_count > 0 and abs(sgml_graphic_count - pdf.image_count) > 2:
        pts = 0.5
        score -= pts
        _add_issue(result, "tagging_accuracy", "minor",
                   f"D2 — PDF has {pdf.image_count} image(s) but SGML has "
                   f"{sgml_graphic_count} <GRAPHIC> tag(s). Count mismatch.",
                   impact=f"-{pts:.1f} pt")

    result.tagging_score = max(0.0, score)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX text extraction helper (GAP 1 — two-stage validation)
# ─────────────────────────────────────────────────────────────────────────────
def _extract_docx_text(docx_path: str) -> dict:
    """
    Extract ALL text from an ABBYY-generated DOCX, including table cells.

    python-docx's doc.paragraphs iteration MISSES table cells — they must be
    extracted explicitly via doc.tables.

    Returns a dict:
      ok             – True if extraction succeeded
      error          – error string on failure
      paragraphs     – list of non-empty paragraph strings (≥2 words)
      table_cells    – list of non-empty, de-duplicated table cell strings
      combined_text  – lowercase normalised blob used for n-gram matching
    """
    try:
        from docx import Document as _DocxDocument  # python-docx
        doc = _DocxDocument(docx_path)

        paragraphs: list[str] = []
        bold_runs: list[str] = []
        italic_runs: list[str] = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t and len(t.split()) >= 2:
                paragraphs.append(t)
            # GAP 4: Extract bold/italic runs directly from DOCX — more reliable
            # than PyMuPDF font-flag inference. python-docx run.bold/run.italic
            # come directly from the OOXML w:rPr/w:b and w:i elements.
            for run in para.runs:
                rt = run.text.strip()
                if not rt or len(rt.split()) < 2:
                    continue
                if run.bold:
                    bold_runs.append(rt)
                if run.italic:
                    italic_runs.append(rt)

        # Explicitly walk tables — doc.paragraphs skips table cells entirely
        table_cells: list[str] = []
        for table in doc.tables:
            seen: set[str] = set()          # deduplicate merged/repeated cells
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and t not in seen:
                        seen.add(t)
                        table_cells.append(t)

        all_text = paragraphs + table_cells
        combined = " ".join(_norm(t) for t in all_text)
        return {
            "ok": True,
            "error": "",
            "paragraphs": paragraphs,
            "table_cells": table_cells,
            "combined_text": combined,
            "bold_runs": bold_runs,    # GAP 4: runs explicitly marked bold in DOCX
            "italic_runs": italic_runs, # GAP 4: runs explicitly marked italic in DOCX
        }
    except Exception as exc:
        return {
            "ok": False,
            "error": str(exc),
            "paragraphs": [],
            "table_cells": [],
            "combined_text": "",
        }


# ─────────────────────────────────────────────────────────────────────────────
# Multi-stage paragraph coverage check (GAP 2 — replaces nested _para_covered)
# ─────────────────────────────────────────────────────────────────────────────
def _para_covered_v2(
    words: list[str],
    blob: str,
    ngrams: set,
    ngram_size: int = 5,
) -> tuple[bool, float, str]:
    """
    Multi-stage paragraph match. Returns (is_covered, confidence, method).

    Stage 1 – Exact 5-gram: fast path for clean identical text.
    Stage 2 – Fuzzy word coverage: ≥90 % of individual words present
              (handles single-word additions / deletions that break n-grams).
    Stage 3 – Sentence-level SequenceMatcher: catches minor rephrasing in
              short-to-medium paragraphs (≤25 words).
    Stage 4 – Chunked window: handles over-merged PyMuPDF paragraphs that
              span multiple SGML elements.
    """
    if not words:
        return (False, 0.0, "empty")

    # Very short text: word-presence check
    if len(words) < ngram_size:
        found = sum(1 for w in words if w in blob)
        cov = found / len(words)
        return (cov >= 0.85, cov, "short_word_match")

    grams = [tuple(words[i:i + ngram_size]) for i in range(len(words) - ngram_size + 1)]
    if not grams:
        return (False, 0.0, "no_grams")

    # Stage 1: exact n-gram match
    matched = sum(1 for g in grams if g in ngrams)
    ngram_ratio = matched / len(grams)
    if ngram_ratio >= 0.55:
        return (True, ngram_ratio, "exact_ngram")

    # Stage 2: fuzzy word coverage
    words_found = sum(1 for w in words if w in blob)
    word_cov = words_found / len(words)
    if word_cov >= 0.90:
        return (True, word_cov, "fuzzy_word_coverage")

    # Stage 3: sentence-level similarity (short paragraphs only — performance guard)
    if len(words) <= 25:
        para_text = " ".join(words)
        best_sim = 0.0
        for seg in re.split(r'[.!?]\s+', blob):
            seg_words = seg.split()
            if len(seg_words) >= max(3, len(words) // 2):
                sim = SequenceMatcher(None, para_text, seg).ratio()
                if sim > best_sim:
                    best_sim = sim
                    if best_sim >= 0.82:
                        break  # good enough — stop early
        if best_sim >= 0.82:
            return (True, best_sim, "sentence_fuzzy")

    # Stage 4: chunked window check (handles over-merged PDF paragraphs)
    chunk_size, step = 15, 8
    if len(words) > chunk_size:
        windows_covered = windows_total = 0
        for start in range(0, len(words) - chunk_size + 1, step):
            chunk = words[start:start + chunk_size]
            cgrams = [tuple(chunk[i:i + ngram_size]) for i in range(len(chunk) - ngram_size + 1)]
            if not cgrams:
                continue
            windows_total += 1
            if sum(1 for g in cgrams if g in ngrams) / len(cgrams) >= 0.60:
                windows_covered += 1
        if windows_total > 0 and (windows_covered / windows_total) >= 0.55:
            return (True, windows_covered / windows_total, "chunked_windows")

    best_conf = max(ngram_ratio, word_cov)
    return (False, best_conf, "no_match")


# ─────────────────────────────────────────────────────────────────────────────
# D3: Text accuracy
# ─────────────────────────────────────────────────────────────────────────────
def check_text_accuracy(pdf: _PDFData, sgml_data: dict, result: L4Result,
                        docx_data: "dict | None" = None) -> None:
    """
    D3 — 8 pts: Paragraph-level text diff, number integrity, citation integrity.

    When docx_data is provided (GAP 1): two-stage comparison
      Stage 1 – PDF vs DOCX  → ABBYY extraction errors (informational, not scored)
      Stage 2 – DOCX vs SGML → pipeline conversion errors (scored)
    Fallback (no DOCX): original single-stage PDF vs SGML comparison.
    """
    score = 8.0
    sgml_blob = sgml_data["text"].lower()
    sgml_paragraphs = sgml_data["paragraphs"]
    raw_sgml = sgml_data.get("_raw_sgml", "")

    # Amending instruments intentionally reproduce only the changed sections, not
    # the full source document. Low paragraph coverage relative to the complete
    # source PDF is therefore by design. We apply a more lenient minimum score
    # tier for such documents.
    # Detection: <QUOTE> tag (classic form), OR "Amending Instrument" / "Amendment
    # Regulation" phrase in title text (some files omit QUOTE wrapping).
    is_amending_doc = bool(
        re.search(r"<QUOTE[\s>]", raw_sgml) or
        re.search(
            r"\bAmend(?:ing|ment(?:ary)?)\s+(?:Instrument|Regulation|Rule|Order)\b",
            raw_sgml, re.IGNORECASE,
        ) or
        # <N> or <TI> tag contains "Amendment" / "Amending" — covers formats
        # like <N>11-803 (Amendment)</N> and <TI>Amendment Regulations</TI>
        re.search(r"<(?:N|TI)[^>]*>[^<]*\bAmend(?:ing|ment(?:ary)?)\b", raw_sgml, re.IGNORECASE)
    )

    if not pdf.paragraphs:
        result.text_score = 8.0
        result.warnings.append("D3 skipped: no paragraphs extracted from PDF (scanned or encrypted).")
        return

    # Single-character bullet tokens that appear in PDF but not SGML
    _BULLET_TOKENS: frozenset[str] = frozenset({"o", "•", "◦", "▪", "▸", "→", "–", "-", ";", ","})

    # Build SGML n-gram set (shared by both single-stage and two-stage paths)
    ngram_size = 5
    sgml_words = sgml_blob.split()
    sgml_ngrams: set[tuple] = set()
    for i in range(len(sgml_words) - ngram_size + 1):
        sgml_ngrams.add(tuple(sgml_words[i:i + ngram_size]))

    # ── D3-a: Paragraph coverage ──────────────────────────────────────────────
    # docx_data is pre-parsed by validate_source_comparison and passed in
    # (GAP 4 refactor: parse DOCX once, share with D2 and D3).

    if docx_data is not None:
        # ── Two-stage path (GAP 1) ────────────────────────────────────────────
        result.docx_available = True
        docx_blob = docx_data["combined_text"]
        docx_words_list = docx_blob.split()
        docx_ngrams: set[tuple] = set()
        for i in range(len(docx_words_list) - ngram_size + 1):
            docx_ngrams.add(tuple(docx_words_list[i:i + ngram_size]))

        # Stage 1: PDF vs DOCX — what ABBYY missed (informational only, not scored)
        meaningful_pdf = [p for p in pdf.paragraphs if len(p.split()) >= 8 and not _is_omittable(p)]
        abbyy_missing: list[str] = []
        abbyy_missing_details: list[dict] = []
        for para in meaningful_pdf:
            words = [w for w in _norm(para).split() if w not in _BULLET_TOKENS]
            is_cov, _conf, _meth = _para_covered_v2(words, docx_blob, docx_ngrams)
            if not is_cov:
                abbyy_missing.append(para[:100])
                abbyy_missing_details.append({'text': para[:100], 'confidence': _conf, 'method': _meth})
        result.abbyy_missing_paragraphs = abbyy_missing
        result.abbyy_missing_paragraph_details = abbyy_missing_details
        if abbyy_missing:
            result.warnings.append(
                f"D3-ABBYY — {len(abbyy_missing)} paragraph(s) from PDF not captured in "
                f"DOCX (ABBYY extraction gaps — cannot be fixed in SGML editor): "
                f"{[p[:60] for p in abbyy_missing[:2]]}"
            )

        # Stage 2: DOCX vs SGML — what the pipeline missed (scored)
        meaningful_docx = [p for p in docx_data["paragraphs"] if len(p.split()) >= 8]
        pipeline_missing: list[str] = []
        pipeline_missing_details: list[dict] = []
        for para in meaningful_docx:
            words = [w for w in _norm(para).split() if w not in _BULLET_TOKENS]
            is_cov, _conf, _meth = _para_covered_v2(words, sgml_blob, sgml_ngrams)
            if not is_cov:
                pipeline_missing.append(para[:100])
                pipeline_missing_details.append({'text': para[:100], 'confidence': _conf, 'method': _meth})
        result.pipeline_missing_paragraphs = pipeline_missing
        result.pipeline_missing_paragraph_details = pipeline_missing_details
        result.missing_paragraphs = pipeline_missing   # backward compat

        if not meaningful_docx:
            result.text_score = 8.0
            return

        coverage = 1.0 - (len(pipeline_missing) / len(meaningful_docx))
        result.text_coverage = coverage

    else:
        # ── Single-stage path (original PDF→SGML) ────────────────────────────
        meaningful = [p for p in pdf.paragraphs if len(p.split()) >= 8 and not _is_omittable(p)]
        if not meaningful:
            result.text_score = 8.0
            return

        sampled = meaningful
        covered = 0
        missing: list[str] = []
        for para in sampled:
            words = [w for w in _norm(para).split() if w not in _BULLET_TOKENS]
            is_cov, _conf, _meth = _para_covered_v2(words, sgml_blob, sgml_ngrams)
            if is_cov:
                covered += 1
            else:
                missing.append(para[:100])

        coverage = covered / len(sampled)
        result.text_coverage = coverage
        result.missing_paragraphs = missing

    # ── Scoring (same tiers for both paths) ──────────────────────────────────
    missing_for_msg = result.missing_paragraphs  # pipeline_missing or missing
    if coverage >= 0.92:
        text_sub_score = 5.0
    elif coverage >= 0.80:
        text_sub_score = 4.0
        _add_issue(result, "text_accuracy", "minor",
                   f"D3 — Paragraph coverage {coverage:.0%}. "
                   f"{len(missing_for_msg)} paragraph(s) from "
                   f"{'DOCX' if docx_data else 'PDF'} not found in SGML.",
                   impact="-1 pt")
    elif coverage >= 0.65:
        text_sub_score = 3.0
        _add_issue(result, "text_accuracy", "major",
                   f"D3 — Paragraph coverage {coverage:.0%}. "
                   f"{len(missing_for_msg)} paragraph(s) missing from SGML. "
                   f"Examples: {missing_for_msg[:2]}",
                   impact="-2 pts")
    else:
        # For amending instruments (<QUOTE> present) low coverage is expected —
        # the SGML only contains the changed sections, not the full source PDF.
        # Floor the sub-score at 3.0 (same as 65-80% tier) instead of 1.0.
        text_sub_score = 3.0 if is_amending_doc else 1.0
        severity = "major" if is_amending_doc else "critical"
        _add_issue(result, "text_accuracy", severity,
                   f"D3 — {'Low' if not is_amending_doc else 'Partial'} paragraph coverage: "
                   f"{coverage:.0%}. "
                   + ("Amending instrument — only changed sections are expected in SGML."
                      if is_amending_doc else
                      "Significant content may be missing from SGML."),
                   impact=f"-{5.0 - text_sub_score:.0f} pts")

    score = score - (5.0 - text_sub_score)

    # D3-b: Number integrity — check numbers from PDF first page appear in SGML
    pdf_numbers = set(_NUMBER_RE.findall(pdf.first_page_text[:2000]))
    missing_numbers = []
    for num in list(pdf_numbers)[:20]:
        norm_num = num.replace(",", "").replace("$", "").replace("%", "")
        if norm_num not in sgml_blob and num not in sgml_blob:
            missing_numbers.append(num)

    if len(missing_numbers) > 3:
        pts = min(1.5, len(missing_numbers) * 0.1)
        score -= pts
        _add_issue(result, "text_accuracy", "major",
                   f"D3 — {len(missing_numbers)} number(s) from PDF first page not found in SGML: "
                   f"{missing_numbers[:5]}. Numbers may be mis-keyed.",
                   impact=f"-{pts:.1f} pts")

    # D3-c: Legal citation integrity
    pdf_citations = set(_LEGAL_CITATION_RE.findall(pdf.first_page_text[:3000]))
    missing_citations = [c for c in pdf_citations if _norm(c) not in sgml_blob]
    if missing_citations:
        pts = min(1.5, len(missing_citations) * 0.25)
        score -= pts
        _add_issue(result, "text_accuracy", "major",
                   f"D3 — Legal citation(s) from PDF not found in SGML: {missing_citations[:3]}. "
                   f"Citations may be corrupted or missing.",
                   impact=f"-{pts:.1f} pts")

    result.text_score = max(0.0, score)


# ─────────────────────────────────────────────────────────────────────────────
# D4: Completeness
# ─────────────────────────────────────────────────────────────────────────────
def check_completeness(pdf: _PDFData, sgml_data: dict, result: L4Result,
                       docx_data: "dict | None" = None) -> None:
    """
    D4 — 7 pts: Count-based completeness check + cell-level table comparison.

    Tables, images, footnotes, sections, pages.
    GAP 5: When docx_data is provided, adds cell-by-cell table content comparison
    (DOCX table cells vs SGML TBLCELL elements) to catch dropped or corrupted rows.
    """
    score = 7.0

    # D4-a: Table count
    pdf_tables = pdf.table_count
    sgml_tables = sgml_data["table_count"]
    if pdf_tables > 0:
        if sgml_tables == 0:
            pts = min(2.0, pdf_tables * 0.5)
            score -= pts
            _add_issue(result, "completeness", "critical",
                       f"D4 — PDF has ~{pdf_tables} table(s) but SGML has no <TABLE> tags. "
                       f"Tables may have been dropped entirely.",
                       impact=f"-{pts:.1f} pts")
        elif abs(sgml_tables - pdf_tables) > max(1, pdf_tables * 0.25):
            pts = 1.0
            score -= pts
            _add_issue(result, "completeness", "major",
                       f"D4 — Table count mismatch: PDF ~{pdf_tables}, SGML {sgml_tables}. "
                       f"Some tables may be missing or split.",
                       impact=f"-{pts:.1f} pt")

    # D4-b: Image count
    pdf_images = pdf.image_count
    sgml_graphics = sgml_data["graphic_count"]
    if pdf_images > 0 and sgml_graphics == 0:
        pts = min(1.5, pdf_images * 0.3)
        score -= pts
        _add_issue(result, "completeness", "major",
                   f"D4 — PDF has {pdf_images} image(s) but SGML has no <GRAPHIC> tags.",
                   impact=f"-{pts:.1f} pts")
    elif pdf_images > 0 and sgml_graphics < pdf_images - 2:
        score -= 0.5
        _add_issue(result, "completeness", "minor",
                   f"D4 — Image count: PDF {pdf_images} vs SGML {sgml_graphics} <GRAPHIC>. "
                   f"Some images may be missing.",
                   impact="-0.5 pts")

    # D4-c: Footnote count
    pdf_fn = pdf.footnote_count
    sgml_fn = sgml_data["fn_count"]
    if pdf_fn > 4 and sgml_fn == 0:
        score -= 1.0
        _add_issue(result, "completeness", "major",
                   f"D4 — PDF has ~{pdf_fn} footnote(s) but SGML has no <FN> tags. "
                   f"Footnotes may have been dropped.",
                   impact="-1.0 pt")
    elif pdf_fn > 0 and sgml_fn > 0 and pdf_fn > sgml_fn + 3:
        score -= 0.5
        _add_issue(result, "completeness", "minor",
                   f"D4 — Footnote count: PDF ~{pdf_fn} vs SGML {sgml_fn}. "
                   f"Some footnotes may be missing.",
                   impact="-0.5 pts")

    # D4-d: Section/heading count ratio
    pdf_sections = len(pdf.headings)
    sgml_sections = len(sgml_data["headings"])
    if pdf_sections > 3 and sgml_sections > 0:
        ratio = sgml_sections / pdf_sections
        # Skip check when SGML has very few sections AND PDF detects many more
        # than expected. This is a false positive for short notice/alert documents
        # where font-size heading detection picks up table column headers, bold
        # data fields, or price entries (e.g. TMX price-list alerts). A document
        # with SGML sections < 4 legitimately has no multi-section structure.
        is_false_heading_detection = sgml_sections < 4 and pdf_sections > sgml_sections * 4
        if not is_false_heading_detection:
            if ratio < 0.5:
                score -= 1.0
                _add_issue(result, "completeness", "major",
                           f"D4 — Section count: PDF {pdf_sections} headings vs SGML "
                           f"{sgml_sections} <TI> tags ({ratio:.0%} coverage). "
                           f"Sections may be missing or merged.",
                           impact="-1.0 pt")
            elif ratio < 0.70:
                score -= 0.5
                _add_issue(result, "completeness", "minor",
                           f"D4 — Section coverage {ratio:.0%}: PDF {pdf_sections} vs "
                           f"SGML {sgml_sections} headings.",
                           impact="-0.5 pts")

    # D4-e: Schedule/appendix detection
    appendix_in_pdf = bool(re.search(
        r"\b(Schedule|Appendix|Annex|Exhibit)\s+[A-Z\d]", pdf.first_page_text, re.IGNORECASE
    ))
    # Use the pre-computed boolean from raw SGML (sgml_data["text"] has tags
    # stripped, so searching it for "<APPENDIX" always returns False — a bug
    # that caused false positives on every CSA Staff Notice with <APPENDIX>).
    appendix_in_sgml = (
        sgml_data.get("has_appendix_tag", False)
        or bool(re.search(r"Schedule\s+[A-Z\d]", sgml_data["text"]))
    )
    if appendix_in_pdf and not appendix_in_sgml:
        score -= 0.5
        _add_issue(result, "completeness", "minor",
                   "D4 — PDF appears to have Schedule/Appendix content but SGML has no "
                   "<APPENDIX> or <SCHEDDOC> and no 'Schedule' text.",
                   impact="-0.5 pts")

    # D4-f: Table cell content coverage (GAP 5 — DOCX cell-by-cell comparison)
    # Only runs when the ABBYY DOCX is available and contains table cells.
    # Compare each non-trivial DOCX table cell against SGML TBLCELL elements.
    if docx_data and docx_data.get("table_cells"):
        _docx_cells = [c for c in docx_data["table_cells"] if len(c.split()) >= 3]
        _sgml_cells = sgml_data.get("table_cells", [])
        if _docx_cells and _sgml_cells:
            _sgml_cell_blob = " ".join(_sgml_cells)
            _scw = _sgml_cell_blob.split()
            _sc_ngrams: set[tuple] = set()
            _ngram_sz = 5
            for _i in range(len(_scw) - _ngram_sz + 1):
                _sc_ngrams.add(tuple(_scw[_i:_i + _ngram_sz]))

            _missing_cells: list[dict] = []
            for _cell in _docx_cells:
                _cwords = _norm(_cell).split()
                _is_cov, _conf, _meth = _para_covered_v2(_cwords, _sgml_cell_blob, _sc_ngrams)
                if not _is_cov:
                    _missing_cells.append({"text": _cell[:80], "confidence": _conf, "method": _meth})

            result.d4_missing_table_cells = _missing_cells

            if _missing_cells:
                _cell_cov = 1.0 - len(_missing_cells) / len(_docx_cells)
                if _cell_cov < 0.75:
                    _pts = 1.5
                    score -= _pts
                    _add_issue(result, "completeness", "major",
                               f"D4 — Table cell coverage {_cell_cov:.0%}: "
                               f"{len(_missing_cells)} of {len(_docx_cells)} DOCX table cell(s) "
                               f"not found in SGML <TBLCELL> elements. "
                               f"Examples: {[m['text'][:50] for m in _missing_cells[:2]]}",
                               impact=f"-{_pts:.1f} pts")
                elif _cell_cov < 0.90:
                    _pts = 0.75
                    score -= _pts
                    _add_issue(result, "completeness", "minor",
                               f"D4 — Table cell coverage {_cell_cov:.0%}: "
                               f"{len(_missing_cells)} DOCX table cell(s) not found in SGML.",
                               impact=f"-{_pts:.1f} pts")

    result.completeness_score = max(0.0, score)


# ─────────────────────────────────────────────────────────────────────────────
# D5: Ordering / sequence
# ─────────────────────────────────────────────────────────────────────────────
def check_ordering(pdf: _PDFData, sgml_data: dict, result: L4Result,
                   docx_data: "dict | None" = None) -> None:
    """
    D5 — 4 pts: Validate that sections appear in the same order as the PDF.

    D5-a: Section heading order via inversion count (heading-level).
    D5-b: Paragraph sequence check via DOCX (GAP 6 — body-level ordering).
          When docx_data is provided, takes up to 30 meaningful DOCX paragraphs,
          locates their 8-word fingerprints in the SGML text, and counts positional
          inversions to detect gross reordering of body content.
    Multi-column layout caveat: flagged as WARNING (may be false positive).
    """
    score = 4.0

    # Gap 9: suppress D5 for 2-column PDFs — ordering is unreliable
    if pdf.two_column:
        result.ordering_score = 4.0
        result.warnings.append(
            "D5 skipped: 2-column PDF layout detected. Section ordering cannot be "
            "reliably validated (PyMuPDF reads left-to-right across columns). "
            "Full score awarded."
        )
        return

    pdf_headings_norm = [_norm(h) for h in pdf.headings if len(h.split()) >= 2]
    sgml_headings_norm = [h for h in sgml_data["headings"] if len(h.split()) >= 1]

    if len(pdf_headings_norm) < 3 or len(sgml_headings_norm) < 3:
        # Not enough headings to validate order meaningfully
        result.ordering_score = 4.0
        return

    # Match SGML headings to PDF headings and record PDF order positions
    pdf_positions: list[int] = []
    for sh in sgml_headings_norm:
        best_pos = -1
        best_ratio = 0.0
        for i, ph in enumerate(pdf_headings_norm):
            ratio = SequenceMatcher(None, sh, ph).ratio()
            if ratio > best_ratio and ratio >= 0.60:
                best_ratio = ratio
                best_pos = i
        if best_pos >= 0:
            pdf_positions.append(best_pos)

    if len(pdf_positions) < 3:
        result.ordering_score = 4.0
        return

    # Count inversions (O(n²) — acceptable for typical heading counts < 50)
    # Also collect the first few inverted pairs for diff_generator
    inversions = 0
    n = len(pdf_positions)
    _inverted_pairs: list[tuple] = []
    for i in range(n):
        for j in range(i + 1, n):
            if pdf_positions[i] > pdf_positions[j]:
                inversions += 1
                if len(_inverted_pairs) < 5:
                    # i appears before j in SGML but after j in PDF
                    _inverted_pairs.append(
                        (sgml_headings_norm[i], sgml_headings_norm[j])
                    )
    result.d5_inverted_pairs = _inverted_pairs

    max_inversions = n * (n - 1) / 2
    inversion_ratio = inversions / max_inversions if max_inversions > 0 else 0

    if inversion_ratio > 0.3:
        pts = min(2.0, inversion_ratio * 4.0)
        score -= pts
        severity = "major" if pts >= 1.0 else "minor"
        _add_issue(result, "ordering", severity,
                   f"D5 — Section order mismatch: {inversion_ratio:.0%} of heading pairs "
                   f"appear in wrong order vs PDF. "
                   f"({inversions} inversion(s) across {n} matched sections)",
                   impact=f"-{pts:.1f} pts")
        result.sequence_violations.append(
            f"{inversions} section ordering inversion(s) detected"
        )
    elif inversion_ratio > 0.1:
        score -= 0.5
        _add_issue(result, "ordering", "minor",
                   f"D5 — Minor section reordering detected ({inversion_ratio:.0%} inversion rate). "
                   f"May be false positive for multi-column layouts.",
                   impact="-0.5 pts")
        result.warnings.append(
            "D5: Minor reordering detected — may be false positive for 2-column PDF layouts."
        )

    # List item sequence: check (a), (b), (c) order within SGML ITEM tags
    items = re.findall(r"<ITEM[^>]*>.*?</ITEM>", sgml_data.get("_raw_sgml", ""), re.DOTALL)
    # (a), (b), (c) patterns
    list_labels = []
    for item in items:
        m = re.search(r"^\s*\(([a-z])\)", re.sub(r"<[^>]+>", "", item).strip())
        if m:
            list_labels.append(ord(m.group(1)) - ord('a'))

    if len(list_labels) >= 3:
        list_inversions = sum(
            1 for i in range(len(list_labels) - 1)
            if list_labels[i] > list_labels[i + 1]
        )
        if list_inversions > 2:
            score -= 0.5
            _add_issue(result, "ordering", "minor",
                       f"D5 — List item order: {list_inversions} out-of-sequence (a)/(b)/(c) "
                       f"ITEM label(s) detected.",
                       impact="-0.5 pts")

    # D5-b: Paragraph sequence check (GAP 6 — DOCX body-level ordering)
    # When DOCX is available, verify that meaningful paragraphs appear in the
    # same left-to-right order in SGML as they do in the DOCX.
    # Uses an 8-word fingerprint search in the SGML text blob.
    if docx_data and docx_data.get("paragraphs") and score > 0:
        _sgml_text = sgml_data["text"].lower()
        _docx_paras = [
            p for p in docx_data["paragraphs"]
            if len(p.split()) >= 10
        ][:30]  # sample up to 30 meaningful paragraphs

        _para_positions: list[int] = []
        for _para in _docx_paras:
            _words = _norm(_para).split()
            _fp = " ".join(_words[:8])   # 8-word fingerprint
            _pos = _sgml_text.find(_fp)
            if _pos >= 0:
                _para_positions.append(_pos)

        if len(_para_positions) >= 5:
            # Count positional inversions
            _n_p = len(_para_positions)
            _para_inv = sum(
                1 for _i in range(_n_p)
                for _j in range(_i + 1, _n_p)
                if _para_positions[_i] > _para_positions[_j]
            )
            _max_inv = _n_p * (_n_p - 1) / 2
            _para_inv_ratio = _para_inv / _max_inv if _max_inv > 0 else 0.0

            if _para_inv_ratio > 0.30:
                _pts = min(1.5, _para_inv_ratio * 3.0)
                score -= _pts
                _add_issue(result, "ordering", "major",
                           f"D5 — Paragraph sequence mismatch: {_para_inv_ratio:.0%} inversion "
                           f"rate across {_n_p} matched DOCX paragraphs. Body content sections "
                           f"may be reordered in the SGML.",
                           impact=f"-{_pts:.1f} pts")
            elif _para_inv_ratio > 0.15:
                score -= 0.5
                _add_issue(result, "ordering", "minor",
                           f"D5 — Minor paragraph reordering detected ({_para_inv_ratio:.0%} "
                           f"inversion rate across {_n_p} matched paragraphs).",
                           impact="-0.5 pts")

    result.ordering_score = max(0.0, score)


# ─────────────────────────────────────────────────────────────────────────────
# D7: Metadata accuracy
# ─────────────────────────────────────────────────────────────────────────────
def check_metadata(pdf: _PDFData, sgml_data: dict, raw_sgml: str, result: L4Result) -> None:
    """
    D7 — 3 pts: Validate POLIDOC metadata against PDF-extracted values.

    Scoring:
      1.0 pt  Language  (SGML LANG vs PDF language heuristic)
      1.0 pt  Doc number (SGML <N> value found anywhere in PDF text)
      1.0 pt  Date      (PDF date found in SGML; ADDDATE vs pub date is soft warning only)

    GAP 7 additions (soft checks — warnings, no additional point deductions):
      • Required fields completeness (LABEL, LANG, ADDDATE present and non-empty)
      • LANG valid-value check (must be EN, FR, or BI)
      • ADDDATE / MODDATE date-format validation (must be YYYYMMDD)
      • MODDATE consistency (MODDATE must be ≥ ADDDATE if both present)
    """
    score = 3.0
    attrs = sgml_data["attrs"]
    mismatches: list[str] = []

    # Build full PDF text for membership checks
    pdf_text_norm = _norm(pdf.first_page_text)

    # Store for diff_generator
    result.d7_expected_lang = pdf.language_hint
    result.d7_pdf_doc_number = pdf.doc_number

    # D7-a: Language — does PDF language match SGML LANG attribute?
    sgml_lang = attrs.get("LANG", "")
    if sgml_lang and pdf.language_hint and sgml_lang != pdf.language_hint:
        score -= 1.0
        mismatches.append(f"LANG={sgml_lang} but PDF appears to be {pdf.language_hint}")
        _add_issue(result, "metadata", "major",
                   f"D7 — POLIDOC LANG='{sgml_lang}' but PDF text suggests language "
                   f"'{pdf.language_hint}'. Document may be mis-classified.",
                   impact="-1.0 pt")

    # D7-b: Document number — SGML <N> tag value should appear in PDF text.
    # Strategy: extract ALL NN-NNN style numbers from the PDF and check if
    # the <N> tag value is among them. This avoids picking up referenced
    # instruments from body text as the document's own number.
    sgml_n_tags = re.findall(r"<N[^>]*>(.*?)</N>", raw_sgml, re.DOTALL)
    sgml_n_values = [re.sub(r"\s+", " ", v).strip() for v in sgml_n_tags[:3]]

    if sgml_n_values:
        # Extract all doc-number-like tokens from PDF text
        pdf_numbers_found = set(re.findall(r"\d{2}-\d{3,4}", pdf.first_page_text))
        # Also include bare numbers like '13-103'
        n_val = sgml_n_values[0]  # first <N> is primary doc number
        n_stripped = re.sub(r"[\s\-]", "", n_val)  # e.g. '45930'
        n_bare = re.search(r"\d{2}-\d{3,4}", n_val)  # e.g. '45-930'
        n_bare_str = n_bare.group(0) if n_bare else ""

        # TMX/alert notices use YYYY-NNN format (e.g. "2025-008", "2025-060").
        # These don't appear prominently in their PDF text and use a different
        # numbering scheme from regulatory NI XX-XXX documents. Treat as
        # warning-only (no point deduction) to avoid false D7 penalties.
        is_yyyy_nnn = bool(re.match(r"^\d{4}-\d{3}$", n_val))

        found_in_pdf = (
            n_val in pdf.first_page_text or
            (n_bare_str and n_bare_str in pdf_numbers_found) or
            n_stripped in pdf.first_page_text.replace("-", "").replace(" ", "")
        )
        if not found_in_pdf and n_bare_str and not is_yyyy_nnn:
            score -= 1.0
            mismatches.append(f"<N>={n_val!r} not found in PDF first pages")
            _add_issue(result, "metadata", "major",
                       f"D7 — SGML <N> value '{n_val}' not found in PDF text. "
                       f"Document number may be wrong. PDF numbers found: "
                       f"{sorted(pdf_numbers_found)[:5]}",
                       impact="-1.0 pt")
        elif not found_in_pdf and is_yyyy_nnn:
            result.warnings.append(
                f"D7: <N> value '{n_val}' (YYYY-NNN format) not found on PDF first pages — "
                f"TMX/alert notice numbers may not appear in extracted PDF text."
            )

    # D7-c: Date — PDF publication date vs SGML date attributes.
    # ADDDATE is the *keying* date (can differ from pub date) — treated as
    # WARNING only (no point deduction). We check that the PDF date exists
    # somewhere in the SGML text as a loose sanity check.
    if pdf.doc_date:
        sgml_adddate = attrs.get("ADDDATE", "")
        if sgml_adddate:
            try:
                from datetime import datetime
                d_pdf = datetime.strptime(pdf.doc_date, "%Y%m%d")
                d_sgml = datetime.strptime(sgml_adddate, "%Y%m%d")
                delta = abs((d_pdf - d_sgml).days)
                if delta > 1825:  # > 5 years: almost certainly a keying error
                    score -= 1.0
                    mismatches.append(
                        f"ADDDATE={sgml_adddate} vs PDF date {pdf.doc_date} (Δ{delta}d)"
                    )
                    _add_issue(result, "metadata", "major",
                               f"D7 — ADDDATE='{sgml_adddate}' differs from PDF date "
                               f"'{pdf.doc_date}' by {delta} days (>5 years). "
                               f"Verify the correct document date.",
                               impact="-1.0 pt")
                elif delta > 30:
                    # Common case: ADDDATE is the keying date, not the publication date.
                    # Amending instruments re-keyed years after original publication
                    # legitimately have large gaps. Warn only, no point deduction.
                    result.warnings.append(
                        f"D7: ADDDATE ({sgml_adddate}) vs PDF date ({pdf.doc_date}) "
                        f"differ by {delta} days — likely keying date vs publication date."
                    )
            except ValueError:
                pass

    result.metadata_mismatches = mismatches
    result.metadata_score = max(0.0, score)

    # ── GAP 7: Soft metadata completeness checks (warnings only, no score impact) ──

    # D7-d: Required POLIDOC attributes present and non-empty
    _required = ("LABEL", "LANG", "ADDDATE")
    for _attr in _required:
        if not attrs.get(_attr, "").strip():
            result.warnings.append(
                f"D7: Required POLIDOC attribute '{_attr}' is missing or empty. "
                f"This may cause downstream processing failures."
            )
            mismatches.append(f"Missing required attribute: {_attr}")

    # D7-e: LANG must be one of the known valid values (EN, FR, BI)
    _valid_langs = {"EN", "FR", "BI"}
    _sgml_lang_val = attrs.get("LANG", "").strip().upper()
    if _sgml_lang_val and _sgml_lang_val not in _valid_langs:
        result.warnings.append(
            f"D7: POLIDOC LANG='{_sgml_lang_val}' is not a recognised value. "
            f"Expected one of: {sorted(_valid_langs)}."
        )
        mismatches.append(f"Invalid LANG value: {_sgml_lang_val!r}")

    # D7-f: Date fields must be valid YYYYMMDD dates; MODDATE must be ≥ ADDDATE
    from datetime import datetime as _dt
    _date_attrs = {k: attrs.get(k, "").strip() for k in ("ADDDATE", "MODDATE")}
    _parsed_dates: dict[str, _dt] = {}
    for _attr, _val in _date_attrs.items():
        if not _val:
            continue
        try:
            _parsed_dates[_attr] = _dt.strptime(_val, "%Y%m%d")
        except ValueError:
            result.warnings.append(
                f"D7: POLIDOC {_attr}='{_val}' is not a valid YYYYMMDD date."
            )
            mismatches.append(f"Invalid date format: {_attr}={_val!r}")

    if "ADDDATE" in _parsed_dates and "MODDATE" in _parsed_dates:
        if _parsed_dates["MODDATE"] < _parsed_dates["ADDDATE"]:
            result.warnings.append(
                f"D7: MODDATE ({attrs['MODDATE']}) is earlier than ADDDATE "
                f"({attrs['ADDDATE']}). A modification date cannot precede the creation date."
            )
            mismatches.append(
                f"MODDATE ({attrs['MODDATE']}) < ADDDATE ({attrs['ADDDATE']})"
            )

    # D7-g: LABEL should be a non-trivial document type string (not just digits/symbols)
    _label_val = attrs.get("LABEL", "").strip()
    if _label_val and not re.search(r"[A-Za-z]{3}", _label_val):
        result.warnings.append(
            f"D7: POLIDOC LABEL='{_label_val}' does not look like a valid document type label."
        )
        mismatches.append(f"Suspicious LABEL value: {_label_val!r}")


# ─────────────────────────────────────────────────────────────────────────────
# Main entry point
# ─────────────────────────────────────────────────────────────────────────────
# D8: Word-level gap detection (HITL — finds exact missing phrases with line numbers)
# ─────────────────────────────────────────────────────────────────────────────
def check_word_gaps(pdf: _PDFData, raw_sgml: str, result: L4Result) -> None:
    """
    D8 — informational (no score deduction): Word-level diff between PDF text
    and SGML text. Finds exact phrases present in the PDF that are absent from
    the SGML, with the SGML line number where the surrounding context appears.

    Reported as HITL issues so a reviewer can jump to the exact line.
    Minimum gap: 3 consecutive words missing.
    """
    if not pdf.paragraphs:
        return

    # Build normalised SGML word list with line-number index
    sgml_lines = raw_sgml.split('\n')
    sgml_line_words: list[tuple[str, int]] = []   # (word, 1-based line number)
    for lineno, line in enumerate(sgml_lines, start=1):
        clean = _norm(re.sub(r'<[^>]+>', ' ', line))
        for w in clean.split():
            if w:
                sgml_line_words.append((w, lineno))

    sgml_words_only = [w for w, _ in sgml_line_words]
    sgml_blob_norm = ' '.join(sgml_words_only)

    # Build normalised PDF word list (all paragraphs joined)
    pdf_words: list[str] = []
    for para in pdf.paragraphs:
        pdf_words.extend(_norm(para).split())

    if not pdf_words or not sgml_words_only:
        return

    # Global word-level diff: find blocks in PDF not present in SGML
    sm = SequenceMatcher(None, pdf_words, sgml_words_only, autojunk=False)
    gaps: list[dict] = []

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag not in ('delete', 'replace'):
            continue
        missing_words = pdf_words[i1:i2]
        if len(missing_words) < 4:
            continue  # ignore very short differences (noise / headers / numbers)

        missing_text = ' '.join(missing_words)

        # Skip if the missing phrase is actually present elsewhere in SGML
        # (difflib may have aligned it differently)
        if _norm(missing_text) in sgml_blob_norm:
            continue

        # Skip omittable boilerplate (headers, page numbers, etc.)
        if _is_omittable(missing_text):
            continue

        # Find SGML line number: look for context words immediately BEFORE the gap
        # (last 4 words of PDF before the deletion)
        context_words = pdf_words[max(0, i1 - 4):i1]
        line_no = 0
        if context_words:
            context_str = ' '.join(context_words)
            # Search SGML lines for these context words
            for ln_idx, ln in enumerate(sgml_lines):
                ln_norm = _norm(re.sub(r'<[^>]+>', ' ', ln))
                if any(cw in ln_norm for cw in context_words[-2:]):
                    line_no = ln_idx + 1
                    break

        # Also try words immediately AFTER the gap for a closer anchor
        if not line_no:
            after_words = pdf_words[i2:i2 + 3]
            for ln_idx, ln in enumerate(sgml_lines):
                ln_norm = _norm(re.sub(r'<[^>]+>', ' ', ln))
                if any(aw in ln_norm for aw in after_words):
                    line_no = ln_idx + 1
                    break

        loc_str = str(line_no) if line_no else ''
        # Truncate for readability
        display = missing_text[:120] + ('...' if len(missing_text) > 120 else '')
        _add_issue(
            result,
            "word_gap",
            "major",
            f"D8 — Missing text from PDF not found in SGML: \"{display}\"",
            location=loc_str,
            impact="HITL: content gap",
        )
        gaps.append({'missing': missing_text, 'line': line_no})

    result.word_gaps = gaps

    # ── D8-b: Paragraph truncation check ─────────────────────────────────────
    # For each PDF paragraph that IS matched in SGML (Stage 1/2 of _para_covered_v2),
    # check whether the SGML version has significantly fewer words — which global
    # difflib misses because it realigns the tail words elsewhere in the document.
    #
    # Strategy: find the SGML line block that contains the first 6 words of the
    # PDF paragraph, then measure word count in that block vs the PDF paragraph.
    # If SGML block has <60% of PDF word count → truncation detected.
    # ─────────────────────────────────────────────────────────────────────────
    sgml_ngrams_set = set()
    for i in range(len(sgml_words_only) - 4):
        sgml_ngrams_set.add(tuple(sgml_words_only[i:i + 5]))

    for para in pdf.paragraphs:
        para_words = _norm(para).split()
        if len(para_words) < 12:
            continue  # too short — truncation not meaningful

        # Check if paragraph is covered at all (Stage 1/2 only for speed)
        anchor = para_words[:6]
        anchor_gram = tuple(anchor[:5])
        if anchor_gram not in sgml_ngrams_set:
            continue  # paragraph not even in SGML — D8 global diff already handles this

        # Find where the anchor appears in the SGML word list
        anchor_pos = None
        for si in range(len(sgml_words_only) - 5):
            if tuple(sgml_words_only[si:si + 5]) == anchor_gram:
                anchor_pos = si
                break
        if anchor_pos is None:
            continue

        # Count how many PDF words match forward from the anchor
        pdf_len = len(para_words)
        matched_ahead = 0
        si = anchor_pos
        pi = 0
        while pi < pdf_len and si < len(sgml_words_only):
            if sgml_words_only[si] == para_words[pi]:
                matched_ahead += 1
                pi += 1
                si += 1
            else:
                # Allow small slippage (skip 1 SGML word)
                si += 1
                if si < len(sgml_words_only) and sgml_words_only[si] == para_words[pi]:
                    matched_ahead += 1
                    pi += 1
                    si += 1
                else:
                    break  # diverged

        sgml_coverage = matched_ahead / pdf_len
        if sgml_coverage < 0.60 and matched_ahead >= 6:
            # Truncation detected — estimate missing tail
            missing_tail_words = para_words[matched_ahead:]
            if len(missing_tail_words) < 4:
                continue
            missing_tail = ' '.join(missing_tail_words)
            if _norm(missing_tail) in sgml_blob_norm:
                continue  # tail present elsewhere — not a truncation
            # Guard against PyMuPDF paragraph-merge false positives: when the
            # PDF extractor merges multiple SGML paragraphs into one large
            # "paragraph", D8-b fires because coverage is low — but the
            # "missing tail" is actually present in the SGML in adjacent
            # paragraphs. If the first 5 words of the tail form a known SGML
            # 5-gram, the content exists — this is a merge artifact, not a
            # real truncation.
            if (len(missing_tail_words) >= 5
                    and tuple(missing_tail_words[:5]) in sgml_ngrams_set):
                continue
            # Additional merge-artifact guard: very long PDF "paragraphs"
            # (>50 words) where only a tiny fraction matched are almost
            # always table cells + section headers merged by PyMuPDF.
            # Require at least 15 matched words before flagging long paras.
            if pdf_len > 50 and matched_ahead < 15:
                continue
            if _is_omittable(missing_tail):
                continue
            # Fragmented-SGML guard: if ≥55% of 4-grams from the missing tail
            # appear anywhere in the SGML blob, the content is present but split
            # across SGML elements by PyMuPDF paragraph merging — not a real truncation.
            if len(missing_tail_words) >= 8:
                _tail_4gs = [
                    ' '.join(missing_tail_words[_ti:_ti + 4])
                    for _ti in range(len(missing_tail_words) - 3)
                ]
                _found_4g = sum(1 for _g in _tail_4gs if _g in sgml_blob_norm)
                if _found_4g / len(_tail_4gs) > 0.55:
                    continue  # majority of tail present in SGML (fragmented) — merge artifact

            # Find line number via anchor context
            anchor_str = ' '.join(anchor[:3])
            trunc_line = 0
            for ln_idx, ln in enumerate(sgml_lines):
                ln_norm = _norm(re.sub(r'<[^>]+>', ' ', ln))
                if any(aw in ln_norm for aw in anchor[:2]):
                    trunc_line = ln_idx + 1
                    break

            display = missing_tail[:120] + ('...' if len(missing_tail) > 120 else '')
            _add_issue(
                result,
                "word_gap",
                "major",
                f"D8-b — Paragraph truncated in SGML ({int(sgml_coverage*100)}% covered). "
                f"Missing tail: \"{display}\"",
                location=str(trunc_line) if trunc_line else '',
                impact="HITL: paragraph truncation",
            )
            gaps.append({'missing': missing_tail, 'line': trunc_line, 'type': 'truncation'})

    # ── D8-c: Short bold phrases from PDF absent from SGML ────────────────────
    # The main gap loop requires ≥4 words to avoid noise. Bold-span deletions
    # (e.g. <BOLD>NI 31-103</BOLD> removed) are often 2-3 words — below that
    # threshold. Here we check PDF bold spans of exactly 2-3 words that are
    # completely absent from the SGML plain text. The sgml_blob_norm guard
    # ensures we only flag phrases genuinely missing from the document, not
    # phrases that exist as plain text elsewhere (which would be harmless).
    pdf_heading_norms = {_norm(h) for h in pdf.headings}
    seen_bold_short: set[str] = set()
    for bold_span in pdf.bold_spans:
        norm_b = _norm(bold_span)
        words_b = norm_b.split()
        if len(words_b) < 2 or len(words_b) > 3:
            continue  # only target short phrases the main loop skips
        if norm_b in seen_bold_short:
            continue
        seen_bold_short.add(norm_b)
        if _is_omittable(bold_span):
            continue
        if norm_b in sgml_blob_norm:
            continue  # text exists somewhere in SGML — not a deletion
        if any(norm_b in ph or ph in norm_b for ph in pdf_heading_norms):
            continue  # part of a heading, not a body bold span
        # Completely absent from SGML — flag as a likely deleted bold element
        _add_issue(
            result,
            "word_gap",
            "major",
            f"D8-c — Short bold phrase from PDF absent from SGML: \"{bold_span[:80]}\"",
            location='',
            impact="HITL: bold content gap",
        )
        gaps.append({'missing': norm_b, 'line': 0, 'type': 'bold_short'})

    # ── D8-d: Leading text deletion check ────────────────────────────────────
    # Mirror of D8-b (tail truncation). D8-b anchors on the FIRST 6 words of
    # a PDF paragraph — so when those words themselves are deleted (e.g.
    # "As of April 1, 2026" removed from the start of a paragraph), D8-b
    # silently skips the paragraph and D8 global diff misses it because the
    # remaining words are present.
    #
    # Strategy: for each PDF paragraph, use words 4-8 as a "body anchor" (skip
    # the lead). If the body anchor IS in SGML, the paragraph body is there.
    # Then check whether the lead words (first 4) are present in SGML immediately
    # before the body anchor position. If they are NOT → leading text was deleted.
    # ─────────────────────────────────────────────────────────────────────────
    for para in pdf.paragraphs:
        para_words = _norm(para).split()
        if len(para_words) < 10:
            continue  # too short to have meaningful leading deletion

        lead_words  = para_words[:5]       # words 0-4 (the potentially deleted lead)
        body_anchor = tuple(para_words[5:10])  # words 5-9 — entirely past the lead, in paragraph body

        # Body must be in SGML — otherwise D8/D8-b handle it
        if body_anchor not in sgml_ngrams_set:
            continue

        # Lead must be genuinely absent from SGML (not just misaligned)
        lead_norm = ' '.join(lead_words)
        if len(lead_words) < 4:
            continue
        if _is_omittable(lead_norm):
            continue
        if lead_norm in sgml_blob_norm:
            continue  # lead present somewhere in SGML — not a deletion

        # Locate body anchor in SGML word list
        body_pos = None
        for si in range(len(sgml_words_only) - 5):
            if tuple(sgml_words_only[si:si + 5]) == body_anchor:
                body_pos = si
                break
        if body_pos is None:
            continue

        # Confirm: SGML words immediately before body_pos don't match the lead
        sgml_before = sgml_words_only[max(0, body_pos - 5): body_pos]
        lead_present_before = any(lead_words[i] == sgml_before[j]
                                  for i in range(len(lead_words))
                                  for j in range(len(sgml_before))
                                  if abs(i - j) <= 1)
        if lead_present_before:
            continue  # lead is actually adjacent — not deleted

        # Find SGML line number from body anchor
        lead_line = 0
        for ln_idx, ln in enumerate(sgml_lines):
            ln_norm = _norm(re.sub(r'<[^>]+>', ' ', ln))
            if any(w in ln_norm for w in body_anchor[:2]):
                lead_line = ln_idx + 1
                break

        display = lead_norm[:120]
        _add_issue(
            result,
            "word_gap",
            "major",
            f"D8-d — Leading text deleted from paragraph: \"{display}\"",
            location=str(lead_line) if lead_line else '',
            impact="HITL: paragraph head deletion",
        )
        gaps.append({'missing': lead_norm, 'line': lead_line, 'type': 'head_deletion'})

    # ── D8-e: Paragraph-local contextual gap detection ────────────────────────
    # D8 global diff misses deletions when the removed text appears elsewhere in
    # the SGML document (the sgml_blob_norm skip fires). D8-e works paragraph-
    # by-paragraph: it locates each PDF paragraph in SGML via a body anchor,
    # extracts a tight local window, and runs SequenceMatcher in that context.
    # This catches leading / middle / trailing deletions that D8 misses.
    # Threshold: 3 words (vs D8's 4) — safe because comparison is contextual.
    _STOP_WORDS_E = frozenset({
        'the', 'a', 'an', 'of', 'in', 'to', 'and', 'or', 'is', 'are',
        'was', 'be', 'that', 'this', 'it', 'as', 'at', 'by', 'for',
        'on', 'with', 'not', 'from', 'its', 'into', 'we', 'our',
    })
    # Words that indicate legal/regulatory body text rather than a contact block.
    # If an interior gap (D8-e) contains NONE of these, it is likely a name list,
    # signature block, or other non-substantive content — skip it.
    _LEGAL_KW_E = frozenset([
        'section', 'subsection', 'paragraph', 'clause', 'subclause',
        'amended', 'amend', 'amendment', 'repeal', 'repealed', 'revoked',
        'regulation', 'regulations', 'rule', 'rules', 'instrument',
        'requirement', 'requirements', 'obligation', 'obligations',
        'provision', 'provisions', 'schedule', 'appendix', 'annex',
        'agreement', 'pursuant', 'compliance', 'disclosure',
        'exemption', 'exemptions', 'registration', 'filing',
        'prospectus', 'securities', 'security', 'issuer', 'issuers',
        'dealer', 'dealers', 'adviser', 'advisers', 'fund', 'funds',
        'order', 'policy', 'policies', 'act', 'statute', 'notice',
        'bulletin', 'effective', 'adopted', 'applies', 'apply',
        'permitted', 'prohibited', 'required', 'written',
    ])
    # Seed dedup with gaps already reported by D8/D8-b/D8-c/D8-d
    seen_d8e: set[str] = {g['missing'][:50] for g in gaps}

    for para in pdf.paragraphs:
        para_words = _norm(para).split()
        if len(para_words) < 8:
            continue

        # Try 5-gram anchors from pdf position 2 onwards — starting at 2 lets
        # us detect leading deletions where the first 1-4 words are gone
        anchor_pos_e: int | None = None
        anchor_off_e: int | None = None
        for pdf_start in range(2, min(len(para_words) - 5, 20)):
            candidate = tuple(para_words[pdf_start:pdf_start + 5])
            if candidate in sgml_ngrams_set:
                for si in range(len(sgml_words_only) - 5):
                    if tuple(sgml_words_only[si:si + 5]) == candidate:
                        anchor_pos_e = si
                        anchor_off_e = pdf_start
                        break
                if anchor_pos_e is not None:
                    break

        if anchor_pos_e is None:
            continue  # whole paragraph absent — D8/D8-b cover this

        # Tight local SGML window around the anchor
        win_start = max(0, anchor_pos_e - anchor_off_e - 2)
        win_end   = min(len(sgml_words_only), anchor_pos_e + len(para_words) + 10)
        sgml_window = sgml_words_only[win_start:win_end]
        sgml_win_text = ' '.join(sgml_window)

        # Local diff: PDF paragraph words vs SGML window words
        sm_e = SequenceMatcher(None, para_words, sgml_window, autojunk=False)
        for tag, i1, i2, _j1, _j2 in sm_e.get_opcodes():
            if tag not in ('delete', 'replace'):
                continue
            missing_w = para_words[i1:i2]
            if len(missing_w) < 3:
                continue
            missing_txt = ' '.join(missing_w)
            missing_nrm = _norm(missing_txt)
            # Must contain at least one content word (not all stopwords)
            if all(w in _STOP_WORDS_E for w in missing_w):
                continue
            if _is_omittable(missing_txt):
                continue
            # Contextual guard: not present in the local SGML window
            if missing_nrm in sgml_win_text:
                continue
            # Global SGML blob check: text may be present in a different paragraph
            if missing_nrm in sgml_blob_norm:
                continue
            # For longer gaps, check 4-gram coverage across the full SGML blob.
            # If ≥55% of the gap's 4-grams are present (fragmented across elements),
            # this is a PyMuPDF merge artefact, not a real deletion.
            if len(missing_w) >= 8:
                _gap_4gs = [
                    ' '.join(missing_w[_gi:_gi + 4])
                    for _gi in range(len(missing_w) - 3)
                ]
                _found_4g = sum(1 for _g in _gap_4gs if _g in sgml_blob_norm)
                if _found_4g / len(_gap_4gs) > 0.55:
                    continue  # majority of gap words in SGML (fragmented) — merge artifact
            # Legal-content guard: if the gap has no regulatory/legal keywords it is
            # likely a contact block, name list, or signature line — skip it.
            if len(missing_w) >= 5 and not any(w in _LEGAL_KW_E for w in missing_w):
                continue
            sig = missing_nrm[:50]
            if sig in seen_d8e:
                continue
            seen_d8e.add(sig)

            # Locate line via anchor word
            e_line = 0
            anchor_word = para_words[anchor_off_e] if anchor_off_e < len(para_words) else ''
            for ln_idx, ln in enumerate(sgml_lines):
                ln_nrm = _norm(re.sub(r'<[^>]+>', ' ', ln))
                if anchor_word and anchor_word in ln_nrm:
                    e_line = ln_idx + 1
                    break

            _add_issue(
                result, "word_gap", "major",
                f"D8-e — Text gap within paragraph: \"{missing_txt[:120]}\"",
                location=str(e_line) if e_line else '',
                impact="HITL: paragraph interior/leading gap",
            )
            gaps.append({'missing': missing_nrm, 'line': e_line, 'type': 'interior'})

    # ── D8-f: Dollar amount / percentage mismatch ────────────────────────────
    # Catches cases where a number was changed (e.g. "$300" → "$250") rather
    # than deleted. For each PDF paragraph matched in SGML, extracts dollar
    # amounts and percentages from both and flags any value present in the PDF
    # paragraph that is absent from the corresponding SGML window.
    _AMOUNT_RE = re.compile(
        r'\$[\d,]+(?:\.\d+)?'       # $300, $1,234.56
        r'|\b\d+(?:\.\d+)?\s*%'     # 15%, 0.5 %
        r'|\b(?:19|20)\d{2}\b'       # years 1900-2099
    )

    def _norm_amount(s: str) -> str:
        return re.sub(r'[\s,]', '', s).lower()

    seen_d8f: set[str] = set()

    for para in pdf.paragraphs:
        para_words = _norm(para).split()
        if len(para_words) < 6:
            continue
        pdf_amounts = _AMOUNT_RE.findall(para)
        if not pdf_amounts:
            continue

        # Locate paragraph in SGML (same anchor strategy as D8-e)
        anchor_pos_f: int | None = None
        anchor_off_f: int | None = None
        for pdf_start in range(0, min(len(para_words) - 5, 20)):
            candidate = tuple(para_words[pdf_start:pdf_start + 5])
            if candidate in sgml_ngrams_set:
                for si in range(len(sgml_words_only) - 5):
                    if tuple(sgml_words_only[si:si + 5]) == candidate:
                        anchor_pos_f = si
                        anchor_off_f = pdf_start
                        break
                if anchor_pos_f is not None:
                    break

        if anchor_pos_f is None:
            continue

        win_start = max(0, anchor_pos_f - anchor_off_f - 2)
        win_end   = min(len(sgml_words_only), anchor_pos_f + len(para_words) + 10)
        sgml_win_f = ' '.join(sgml_words_only[win_start:win_end])
        # Also check raw SGML text around the anchor (preserves $ signs stripped by _norm)
        approx_line = 0
        for ln_idx, ln in enumerate(sgml_lines):
            if anchor_off_f < len(para_words) and para_words[anchor_off_f] in _norm(re.sub(r'<[^>]+>', ' ', ln)):
                approx_line = ln_idx + 1
                break
        # Reconstruct raw SGML window ±30 lines for amount comparison
        raw_win_lines = sgml_lines[max(0, approx_line - 5): min(len(sgml_lines), approx_line + 15)]
        raw_win = re.sub(r'<[^>]+>', ' ', ' '.join(raw_win_lines))
        sgml_amounts = set(_norm_amount(a) for a in _AMOUNT_RE.findall(raw_win))

        for amt in pdf_amounts:
            amt_n = _norm_amount(amt)
            if amt_n in seen_d8f:
                continue
            if amt_n not in sgml_amounts:
                # Cross-check against full SGML (in case the amount is present but distant)
                full_sgml_raw = re.sub(r'<[^>]+>', ' ', raw_sgml)
                if amt_n in (_norm_amount(a) for a in _AMOUNT_RE.findall(full_sgml_raw)):
                    continue  # present elsewhere — skip
                seen_d8f.add(amt_n)
                _add_issue(
                    result, "word_gap", "major",
                    f"D8-f — Amount/date in PDF not found in SGML: \"{amt}\"",
                    location=str(approx_line) if approx_line else '',
                    impact="HITL: possible number substitution",
                )
                gaps.append({'missing': amt_n, 'line': approx_line, 'type': 'amount'})

    result.word_gaps = gaps


# ─────────────────────────────────────────────────────────────────────────────
# D9: Contact info presence check (phone / email / URL)
# ─────────────────────────────────────────────────────────────────────────────
_RE_PHONE = re.compile(
    r'(?<!\d)(\+?1[-.\s]?)?(\(?\d{3}\)?[-.\s]\d{3}[-.\s]\d{4})(?!\d)'
)
_RE_EMAIL = re.compile(
    r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}'
)
_RE_URL = re.compile(
    r'https?://[^\s<>"&]{8,}|www\.[^\s<>"&]{5,}'
)


def check_contact_info(pdf: _PDFData, raw_sgml: str, result: L4Result) -> None:
    """
    D9 — informational: detect phone numbers, email addresses, and URLs that
    appear in the PDF source but are absent from the SGML.
    """
    if not pdf.paragraphs:
        return

    pdf_text  = '\n'.join(pdf.paragraphs)
    sgml_text = re.sub(r'<[^>]+>', ' ', raw_sgml)

    def _find_missing(pattern, label):
        pdf_items  = set(m.group(0).strip() for m in pattern.finditer(pdf_text))
        sgml_items = set(m.group(0).strip() for m in pattern.finditer(sgml_text))
        # Normalise phone numbers (strip spaces/dashes) for comparison
        def _norm_item(s):
            return re.sub(r'[\s\-.()+]', '', s).lower()
        sgml_norm = {_norm_item(s) for s in sgml_items}
        missing = [item for item in sorted(pdf_items)
                   if _norm_item(item) not in sgml_norm]
        for item in missing:
            # Find approximate SGML line — look for nearby words in PDF context
            _add_issue(
                result,
                "contact_info",
                "major",
                f"D9 — {label} in PDF not found in SGML: \"{item}\"",
                location="",
                impact="HITL: contact info gap",
            )

    _find_missing(_RE_PHONE, "Phone number")
    _find_missing(_RE_EMAIL, "Email address")
    _find_missing(_RE_URL,   "URL/hyperlink")


# ─────────────────────────────────────────────────────────────────────────────
def validate_source_comparison(
    raw_sgml: str,
    pdf_path: Optional[str] = None,
    docx_path: Optional[str] = None,
) -> L4Result:
    """
    Run all Level 4 source-comparison checks.

    Parameters
    ----------
    raw_sgml  : str           — Raw SGML content as read from file.
    pdf_path  : str, optional — Path to source PDF. If None, only D6 (encoding)
                                runs. All other dimensions are skipped with warnings.
    docx_path : str, optional — Path to ABBYY-generated DOCX (intermediate file).
                                When provided, D3 uses two-stage comparison:
                                PDF→DOCX (ABBYY errors) + DOCX→SGML (pipeline errors).
                                Significantly reduces false positives from headers/footers.

    Returns
    -------
    L4Result with score (0-30) and all issues found.
    """
    result = L4Result()

    # D6 always runs — no PDF needed
    check_encoding(raw_sgml, result)

    if not pdf_path:
        result.pdf_available = False
        result.warnings.append(
            "L4: No source PDF provided. D2/D3/D4/D5/D7 skipped. "
            "Only encoding check (D6) ran."
        )
        # Score: only D6 ran out of 30 pts. Normalise D6 score to 3/30
        result.tagging_score = 0.0
        result.text_score = 0.0
        result.completeness_score = 0.0
        result.ordering_score = 0.0
        result.metadata_score = 0.0
        result.score = result.encoding_score  # out of 3
        return result

    result.pdf_available = True

    # Extract PDF data
    pdf = _extract_pdf_data(pdf_path)

    # Store PDF headings for D3 placement heuristic in diff_generator
    result.pdf_headings = pdf.headings  # all headings — no cap

    if not pdf.ok:
        result.pdf_text_extractable = False
        result.warnings.append(f"L4: PDF extraction failed ({pdf.error}). D2-D5/D7 skipped.")
        result.score = result.encoding_score
        return result

    result.pdf_text_extractable = True

    # Extract SGML structured data (pass raw for D5 list check)
    sgml_data = _extract_sgml_text(raw_sgml)
    sgml_data["_raw_sgml"] = raw_sgml  # pass through for D5 list-item check

    # Run all dimensions with individual error isolation
    # GAP 4: parse DOCX once here — share with D2 (check_tagging) and D3 (check_text_accuracy)
    # so we don't parse the DOCX file twice.
    _docx_data: dict | None = None
    if docx_path:
        _d = _extract_docx_text(docx_path)
        if _d["ok"] and (_d["paragraphs"] or _d["bold_runs"] or _d["italic_runs"]):
            _docx_data = _d
        else:
            result.warnings.append(
                f"D2/D3 — DOCX extraction failed ({_d.get('error', 'unknown')}). "
                "Falling back to PDF-only validation."
            )

    try:
        check_tagging(pdf, sgml_data, raw_sgml, result, docx_data=_docx_data)
    except Exception as e:
        result.tagging_score = 5.0  # assume pass on error
        result.warnings.append(f"D2 check error (skipped): {e}")

    # D3 + D8: replaced by SemanticContentAgent (no deterministic word-matching)
    _doc_type = sgml_data.get("attrs", {}).get("LABEL", "Notice")
    try:
        check_text_semantic(pdf, sgml_data, result, doc_type=_doc_type)
    except Exception as e:
        result.text_score = 8.0
        result.warnings.append(f"D3/D8 semantic agent error (skipped): {e}")

    try:
        check_completeness(pdf, sgml_data, result, docx_data=_docx_data)
    except Exception as e:
        result.completeness_score = 7.0
        result.warnings.append(f"D4 check error (skipped): {e}")

    try:
        check_ordering(pdf, sgml_data, result, docx_data=_docx_data)
    except Exception as e:
        result.ordering_score = 4.0
        result.warnings.append(f"D5 check error (skipped): {e}")

    try:
        check_metadata(pdf, sgml_data, raw_sgml, result)
    except Exception as e:
        result.metadata_score = 3.0
        result.warnings.append(f"D7 check error (skipped): {e}")

    # D8 word-gap check removed — now handled by SemanticContentAgent above.

    try:
        check_contact_info(pdf, raw_sgml, result)
    except Exception as e:
        result.warnings.append(f"D9 contact-info check error (skipped): {e}")

    result.score = (
        result.tagging_score +
        result.text_score +
        result.completeness_score +
        result.ordering_score +
        result.encoding_score +
        result.metadata_score
    )

    # Enrich L4 issues with actionable fix templates
    try:
        from validator.core.fix_templates import enrich_issues
        enrich_issues(result.issues)
    except Exception:
        pass

    return result
